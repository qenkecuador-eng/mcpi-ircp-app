from pathlib import Path
from io import BytesIO
from datetime import datetime
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt

try:
    from docx import Document
except Exception:
    Document = None

APP_DIR = Path(__file__).parent
BASE_PATH = APP_DIR / "tipos_penales_base_ampliada.csv"
SUBCRITERIOS_PATH = APP_DIR / "subcriterios_MCPI_IRCP_app_unico.xlsx"
SUPUESTOS_PATH = APP_DIR / "supuestos_hecho_base.csv"
TIPICIDAD_PATH = APP_DIR / "tipicidad_por_delito.csv"
ESTANDARES_PATH = APP_DIR / "estandares_convencionalidad.csv"
JURIS_PATH = APP_DIR / "jurisprudencia_base.csv"
CATALOGO_DESC_PATH = APP_DIR / "catalogo_descripciones_tipo_penal.csv"

st.set_page_config(
    page_title="MCPI-IRCP-I v8",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

NUMERIC_COLS = [
    "Legalidad", "Claridad normativa", "Lesividad", "Mínima intervención penal",
    "Idoneidad", "Necesidad penal", "Peso derechos indígenas", "Peso interés estatal"
]

BASE_COLS = [
    "País", "Norma", "Artículo", "Tipo penal", "Categoría", "Estado",
    "Contexto típico de protesta", *NUMERIC_COLS, "Observación metodológica"
]

def load_csv(path, default_cols=None):
    if Path(path).exists():
        return pd.read_csv(path).fillna("")
    return pd.DataFrame(columns=default_cols or [])

def load_base():
    df = load_csv(BASE_PATH, BASE_COLS)
    for c in BASE_COLS:
        if c not in df.columns:
            df[c] = 0 if c in NUMERIC_COLS else ""
    for c in NUMERIC_COLS:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).clip(0, 100)
    return df[BASE_COLS]

def load_rubrica():
    if SUBCRITERIOS_PATH.exists():
        return pd.read_excel(SUBCRITERIOS_PATH, sheet_name=0).fillna("")
    return pd.DataFrame()

def nivel_riesgo(valor):
    v = float(valor)
    if v >= 70:
        return "Riesgo alto", "Rojo"
    if v >= 50:
        return "Riesgo medio-alto", "Naranja"
    if v >= 35:
        return "Riesgo medio", "Amarillo"
    return "Riesgo bajo", "Verde"

def nivel_compat(valor):
    v = float(valor)
    if v >= 75:
        return "Compatibilidad alta", "Verde"
    if v >= 60:
        return "Compatibilidad condicionada", "Amarillo"
    if v >= 40:
        return "Compatibilidad débil", "Naranja"
    return "Incompatibilidad o compatibilidad crítica", "Rojo"

def calcular_indice(valores):
    legalidad = float(valores.get("Legalidad", 0))
    claridad = float(valores.get("Claridad normativa", 0))
    lesividad = float(valores.get("Lesividad", 0))
    minima = float(valores.get("Mínima intervención penal", 0))
    idoneidad = float(valores.get("Idoneidad", 0))
    necesidad = float(valores.get("Necesidad penal", 0))
    pdi = float(valores.get("Peso derechos indígenas", 0))
    pie = float(valores.get("Peso interés estatal", 0))

    formal = round(legalidad * 0.60 + claridad * 0.40, 2)
    material = round(lesividad * 0.60 + minima * 0.40, 2)
    europeo = round(idoneidad * 0.50 + necesidad * 0.50, 2)
    alexy = round((pie / (pdi + pie)) * 100, 2) if (pdi + pie) else 0
    peso_relativo = round(pdi / pie, 2) if pie else 999
    cpi = round(formal * 0.25 + material * 0.30 + europeo * 0.25 + alexy * 0.20, 2)
    ircp = round(100 - cpi, 2)
    nivel, semaforo = nivel_riesgo(ircp)

    return {
        "Legalidad": legalidad,
        "Claridad normativa": claridad,
        "Lesividad": lesividad,
        "Mínima intervención penal": minima,
        "Idoneidad": idoneidad,
        "Necesidad penal": necesidad,
        "Peso derechos indígenas": pdi,
        "Peso interés estatal": pie,
        "Formal": formal,
        "Material": material,
        "Europeo": europeo,
        "Alexy": alexy,
        "Peso relativo": peso_relativo,
        "CPI": cpi,
        "IRCP-I": ircp,
        "Nivel": nivel,
        "Semáforo": semaforo,
    }

def buscar_valores_formales(base, pais, articulo, tipo_penal):
    articulo_simple = str(articulo).replace("art. ", "").replace(" CP", "").replace(" COIP", "")
    mask = (
        base["País"].astype(str).str.lower().eq(str(pais).lower()) &
        (
            base["Artículo"].astype(str).str.contains(articulo_simple, case=False, na=False) |
            base["Tipo penal"].astype(str).str.lower().eq(str(tipo_penal).lower())
        )
    )
    if mask.any():
        r = base[mask].iloc[0]
        return {"Legalidad": float(r["Legalidad"]), "Claridad normativa": float(r["Claridad normativa"])}
    return {"Legalidad": 55.0, "Claridad normativa": 50.0}

def yes(v):
    return str(v).strip().lower() in ["sí", "si", "yes", "true", "1"]

def is_no_or_missing(v):
    return str(v).strip().lower() in ["no", "no consta", "no aplica", "no / deficiente", ""]

def score_desde_hechos(hechos, base):
    formales = buscar_valores_formales(base, hechos.get("pais",""), hechos.get("articulo",""), hechos.get("tipo_penal",""))

    violencia = yes(hechos.get("violencia_grave", "No"))
    danos = yes(hechos.get("danos_personas", "No"))
    destruccion = yes(hechos.get("destruccion_bienes", "No"))
    armas = yes(hechos.get("uso_armas", "No"))
    org = yes(hechos.get("organizacion_criminal", "No"))
    riesgo = yes(hechos.get("riesgo_concreto_vida_integridad", "No"))
    alternativas = yes(hechos.get("alternativas_no_penales", "Sí"))
    consulta_def = "no" in str(hechos.get("consulta_previa", "")).lower() or "deficiente" in str(hechos.get("consulta_previa", "")).lower()
    indigena = "indígena" in str(hechos.get("sujeto_colectivo", "")).lower() or "indigena" in str(hechos.get("sujeto_colectivo", "")).lower()
    extractivo = "extract" in str(hechos.get("contexto", "")).lower()
    parcial = "parcial" in str(hechos.get("nivel_afectacion", "")).lower()
    total = "total" in str(hechos.get("nivel_afectacion", "")).lower()
    grave_servicio = "grave" in str(hechos.get("nivel_afectacion", "")).lower()

    les_dano = 35 if parcial else 45
    les_gravedad = 25
    les_riesgo = 20
    les_prueba = 35 if str(hechos.get("prueba_individualizada","")).lower() in ["no", "no consta", ""] else 60

    if total:
        les_dano += 15
        les_gravedad += 10
    if grave_servicio:
        les_dano += 25
        les_gravedad += 25
    if violencia:
        les_gravedad += 25
    if danos:
        les_dano += 25
        les_gravedad += 20
    if destruccion:
        les_dano += 15
        les_gravedad += 10
    if armas:
        les_gravedad += 15
        les_riesgo += 20
    if riesgo:
        les_riesgo = 70
        les_gravedad += 20
    if org:
        les_prueba += 10

    les_dano, les_gravedad, les_riesgo, les_prueba = [min(100, x) for x in [les_dano, les_gravedad, les_riesgo, les_prueba]]

    min_dialogo = 20 if "no" in str(hechos.get("dialogo_previo","")).lower() or "consta" in str(hechos.get("dialogo_previo","")).lower() else 60
    min_consulta = 15 if consulta_def and indigena else 50
    min_medidas = 30 if alternativas else 65
    min_ultima = 20 if alternativas else 65

    ido_finalidad = 60
    ido_conexion = 35
    ido_riesgo = 20 if not riesgo else 70
    ido_violencia = 10 if not (violencia or danos or destruccion or armas) else 55
    ido_delimitacion = 30
    ido_evidencia = 40 if les_prueba < 50 else 60
    ido_intercultural = 15 if indigena else 35

    if grave_servicio or total:
        ido_conexion += 15

    nec_dialogo = 15 if min_dialogo <= 20 else 55
    nec_consulta = 10 if consulta_def and indigena else 50
    nec_mediacion = 20 if indigena else 35
    nec_admin = 30 if alternativas else 65
    nec_preventivas = 35 if alternativas else 60
    nec_eficacia = 35 if alternativas else 60
    nec_no_agotamiento = 15 if alternativas else 55

    pdi_territorial = 90 if indigena and extractivo else (75 if indigena else 35)
    pdi_abstracto = 90 if indigena else 50
    pdi_fiabilidad = 75 if indigena else 40

    pie_intensidad = 35 if parcial else 50
    if total:
        pie_intensidad = 55
    if grave_servicio:
        pie_intensidad = 70
    if violencia or danos or destruccion or armas or riesgo:
        pie_intensidad += 25
    pie_abstracto = 55
    pie_fiabilidad = 30
    if riesgo or danos:
        pie_fiabilidad += 25

    rows = [
        ["Control formal", "Legalidad", "LEG-BASE", "Valor abstracto del tipo penal seleccionado", 1.00, formales["Legalidad"], "Control formal del tipo penal en abstracto."],
        ["Control formal", "Claridad normativa", "CLA-BASE", "Valor abstracto del tipo penal seleccionado", 1.00, formales["Claridad normativa"], "Claridad normativa del texto penal."],
        ["Control material", "Lesividad", "LES-01", "Daño real", 0.30, les_dano, "Afectación fáctica acreditada."],
        ["Control material", "Lesividad", "LES-02", "Gravedad del daño", 0.25, les_gravedad, "Violencia, daño, armas o gravedad del hecho."],
        ["Control material", "Lesividad", "LES-03", "Concreción del riesgo", 0.25, les_riesgo, "Riesgo concreto para terceros."],
        ["Control material", "Lesividad", "LES-04", "Prueba individualizada", 0.20, les_prueba, "Individualización de autoría, daño y nexo causal."],
        ["Control material", "Mínima intervención penal", "MIN-01", "Agotamiento de diálogo", 0.25, min_dialogo, "Verifica si hubo diálogo previo."],
        ["Control material", "Mínima intervención penal", "MIN-02", "Consulta previa o vía institucional", 0.25, min_consulta, "Consulta previa o canal institucional adecuado."],
        ["Control material", "Mínima intervención penal", "MIN-03", "Existencia de medidas no penales", 0.25, min_medidas, "Alternativas administrativas, civiles o de mediación."],
        ["Control material", "Mínima intervención penal", "MIN-04", "Derecho penal como última ratio", 0.25, min_ultima, "Excepcionalidad de la vía penal."],
        ["Test europeo", "Idoneidad", "IDO-01", "Finalidad legítima", 0.15, ido_finalidad, "Finalidad estatal legítima."],
        ["Test europeo", "Idoneidad", "IDO-02", "Conexión medio-fin", 0.15, ido_conexion, "Conexión entre vía penal y protección del bien jurídico."],
        ["Test europeo", "Idoneidad", "IDO-03", "Riesgo real actual individualizado", 0.20, ido_riesgo, "Riesgo actual y atribuible."],
        ["Test europeo", "Idoneidad", "IDO-04", "Violencia grave o lesividad relevante", 0.20, ido_violencia, "Violencia o daño relevante."],
        ["Test europeo", "Idoneidad", "IDO-05", "Delimitación estricta de aplicación", 0.10, ido_delimitacion, "Diferencia protesta protegida de delito."],
        ["Test europeo", "Idoneidad", "IDO-06", "Evidencia fáctica suficiente", 0.10, ido_evidencia, "Suficiencia probatoria."],
        ["Test europeo", "Idoneidad", "IDO-07", "Enfoque intercultural inicial", 0.10, ido_intercultural, "Consideración inicial de contexto indígena."],
        ["Test europeo", "Necesidad penal", "NEC-01", "Diálogo previo", 0.15, nec_dialogo, "Agotamiento de diálogo."],
        ["Test europeo", "Necesidad penal", "NEC-02", "Consulta previa adecuada", 0.20, nec_consulta, "Consulta previa libre e informada."],
        ["Test europeo", "Necesidad penal", "NEC-03", "Mediación intercultural", 0.15, nec_mediacion, "Mediación antes de penalización."],
        ["Test europeo", "Necesidad penal", "NEC-04", "Medidas administrativas o civiles", 0.15, nec_admin, "Opciones no penales."],
        ["Test europeo", "Necesidad penal", "NEC-05", "Medidas preventivas no penales", 0.10, nec_preventivas, "Prevención y gestión pública del conflicto."],
        ["Test europeo", "Necesidad penal", "NEC-06", "Eficacia de alternativas menos lesivas", 0.15, nec_eficacia, "Eficacia de medidas alternativas."],
        ["Test europeo", "Necesidad penal", "NEC-07", "No agotamiento real de alternativas", 0.10, nec_no_agotamiento, "Prueba de agotamiento estatal."],
        ["Alexy", "Peso derechos indígenas", "PDI-01", "Intensidad de afectación territorial/cultural", 0.35, pdi_territorial, "Territorio, cultura, consulta y ambiente."],
        ["Alexy", "Peso derechos indígenas", "PDI-02", "Peso abstracto del derecho indígena comprometido", 0.35, pdi_abstracto, "Peso jurídico de territorio, consulta, expresión, reunión y participación."],
        ["Alexy", "Peso derechos indígenas", "PDI-03", "Fiabilidad probatoria del contexto indígena", 0.30, pdi_fiabilidad, "Respaldo del contexto indígena."],
        ["Alexy", "Peso interés estatal", "PIE-01", "Intensidad de afectación al interés estatal", 0.35, min(100, pie_intensidad), "Afectación a circulación, seguridad, servicios u orden público."],
        ["Alexy", "Peso interés estatal", "PIE-02", "Peso abstracto del interés estatal", 0.35, pie_abstracto, "Importancia general del bien estatal."],
        ["Alexy", "Peso interés estatal", "PIE-03", "Fiabilidad probatoria del daño estatal", 0.30, min(100, pie_fiabilidad), "Prueba del daño o riesgo estatal."],
    ]

    sub = pd.DataFrame(rows, columns=["Modulo", "Variable", "Codigo", "Subcriterio", "Peso", "Valor", "Explicación metodológica"])
    sub["Aporte"] = (sub["Peso"] * sub["Valor"]).round(2)
    vals = sub.groupby("Variable", as_index=False)["Aporte"].sum().rename(columns={"Aporte": "Valor"})
    valores = {r["Variable"]: float(r["Valor"]) for _, r in vals.iterrows()}

    traz = pd.DataFrame([
        ["Hecho", "Estado", "Variable afectada", "Impacto"],
        ["Tipo penal seleccionado", hechos.get("articulo",""), "Control formal", "Activa valores abstractos de legalidad y claridad."],
        ["Comunidad indígena", hechos.get("sujeto_colectivo",""), "Alexy / convencionalidad", "Eleva peso de derechos indígenas si corresponde."],
        ["Consulta previa", hechos.get("consulta_previa",""), "Mínima intervención / Alexy", "Si es inexistente o deficiente, baja necesidad penal y sube derechos indígenas."],
        ["Violencia grave", hechos.get("violencia_grave",""), "Lesividad / idoneidad", "Si no se acredita, baja compatibilidad de la vía penal."],
        ["Daños a personas", hechos.get("danos_personas",""), "Lesividad", "Si no se acredita, reduce daño penal relevante."],
        ["Uso de armas", hechos.get("uso_armas",""), "Riesgo concreto", "Si no se acredita, reduce riesgo grave."],
        ["Organización criminal", hechos.get("organizacion_criminal",""), "Tipicidad / prueba", "Evita asimilar organización comunitaria a criminalidad."],
        ["Alternativas no penales", hechos.get("alternativas_no_penales",""), "Necesidad penal", "Si existían, la vía penal exige motivación reforzada."],
    ], columns=["Hecho del caso","Estado","Variable afectada","Impacto metodológico"])
    return sub, traz, valores


def validar_pesos_subcriterios(df):
    if df is None or df.empty:
        return pd.DataFrame(columns=["Variable", "Suma de pesos", "Estado"])
    tmp = df.copy()
    tmp["Peso"] = pd.to_numeric(tmp["Peso"], errors="coerce").fillna(0)
    resumen = tmp.groupby("Variable", as_index=False)["Peso"].sum().rename(columns={"Peso": "Suma de pesos"})
    resumen["Suma de pesos"] = resumen["Suma de pesos"].round(4)
    resumen["Estado"] = resumen["Suma de pesos"].apply(lambda x: "Correcto" if abs(float(x) - 1.0) <= 0.001 else "Revisar")
    return resumen

def normalizar_pesos_subcriterios(df):
    if df is None or df.empty:
        return df
    tmp = df.copy()
    tmp["Peso"] = pd.to_numeric(tmp["Peso"], errors="coerce").fillna(0)
    tmp["Valor"] = pd.to_numeric(tmp["Valor"], errors="coerce").fillna(0).clip(0, 100)
    for variable in tmp["Variable"].dropna().unique():
        mask = tmp["Variable"] == variable
        suma = tmp.loc[mask, "Peso"].sum()
        if suma > 0:
            tmp.loc[mask, "Peso"] = tmp.loc[mask, "Peso"] / suma
    tmp["Peso"] = tmp["Peso"].round(6)
    tmp["Aporte"] = (tmp["Peso"] * tmp["Valor"]).round(2)
    return tmp

def calcular_variables_desde_subcriterios(df):
    tmp = df.copy()
    tmp["Peso"] = pd.to_numeric(tmp["Peso"], errors="coerce").fillna(0)
    tmp["Valor"] = pd.to_numeric(tmp["Valor"], errors="coerce").fillna(0).clip(0, 100)
    tmp["Aporte"] = (tmp["Peso"] * tmp["Valor"]).round(2)
    var_df = tmp.groupby("Variable", as_index=False)["Aporte"].sum().rename(columns={"Aporte": "Valor"})
    valores = {r["Variable"]: float(r["Valor"]) for _, r in var_df.iterrows()}
    return tmp, var_df, valores

def componentes(res):
    return pd.DataFrame([
        {"Componente": "Control formal", "Valor": res["Formal"], "Peso CPI": 0.25, "Aporte": round(res["Formal"] * 0.25, 2)},
        {"Componente": "Control material", "Valor": res["Material"], "Peso CPI": 0.30, "Aporte": round(res["Material"] * 0.30, 2)},
        {"Componente": "Test europeo", "Valor": res["Europeo"], "Peso CPI": 0.25, "Aporte": round(res["Europeo"] * 0.25, 2)},
        {"Componente": "Puntaje Alexy", "Valor": res["Alexy"], "Peso CPI": 0.20, "Aporte": round(res["Alexy"] * 0.20, 2)},
    ])

def reglas_cierre(res):
    reglas = []
    if res["Legalidad"] < 40: reglas.append("Legalidad menor a 40: alerta de incompatibilidad formal.")
    if res["Lesividad"] < 40: reglas.append("Lesividad menor a 40: no se acredita daño penal grave suficiente.")
    if res["Mínima intervención penal"] < 40: reglas.append("Mínima intervención menor a 40: debieron priorizarse vías no penales.")
    if res["Idoneidad"] < 75: reglas.append("Idoneidad menor a 75: criminalización penal no plenamente justificada.")
    if res["Necesidad penal"] < 60: reglas.append("Necesidad penal menor a 60: existen o debieron agotarse alternativas menos lesivas.")
    if res["Peso relativo"] > 1: reglas.append("Alexy favorece derechos indígenas/protesta: exige motivación reforzada.")
    if res["IRCP-I"] >= 50: reglas.append(f"IRCP-I {res['IRCP-I']}: riesgo relevante de criminalización.")
    return reglas

def to_excel(sheets):
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        for name, df in sheets.items():
            df.to_excel(writer, index=False, sheet_name=str(name)[:31])
    bio.seek(0)
    return bio.getvalue()

def conclusion(hechos, res):
    return (
        f"La evaluación MCPI-IRCP-I del caso {hechos.get('nombre_caso','')} arroja una CPI de {res['CPI']}/100 "
        f"y un IRCP-I de {res['IRCP-I']}/100, equivalente a {res['Nivel']} / semáforo {res['Semáforo']}. "
        f"El resultado no decide el caso ni reemplaza la valoración judicial, sino que estructura el control de legalidad, "
        f"lesividad, proporcionalidad, mínima intervención y ponderación intercultural. La decisión final exige prueba "
        f"individualizada, contradicción procesal, motivación reforzada y control de convencionalidad."
    )

def docx_report(hechos, res, sub, traz, prob, tip, conv, contrad):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("Informe judicial asistido MCPI-IRCP-I", 0)
    doc.add_paragraph("Herramienta de apoyo. No reemplaza decisión judicial.")
    doc.add_heading("1. Expediente y supuesto", level=1)
    for k in ["nombre_caso","pais","norma","articulo","tipo_penal","sujeto_colectivo","tipo_protesta","finalidad","contexto"]:
        doc.add_paragraph(f"{k}: {hechos.get(k,'')}")
    doc.add_paragraph(str(hechos.get("descripcion","")))
    doc.add_heading("2. Resultado", level=1)
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Indicador"; tbl.rows[0].cells[1].text = "Valor"
    for k in ["Formal","Material","Europeo","Alexy","Peso relativo","CPI","IRCP-I","Nivel","Semáforo"]:
        cells = tbl.add_row().cells; cells[0].text = k; cells[1].text = str(res.get(k,""))
    doc.add_heading("3. Conclusión", level=1)
    doc.add_paragraph(conclusion(hechos, res))
    doc.add_heading("4. Reglas de cierre", level=1)
    for r in reglas_cierre(res): doc.add_paragraph(r)
    doc.add_heading("5. Trazabilidad", level=1)
    for _, r in traz.iterrows():
        doc.add_paragraph(f"{r['Hecho del caso']} | {r['Estado']} | {r['Variable afectada']} | {r['Impacto metodológico']}")
    bio = BytesIO(); doc.save(bio); return bio.getvalue()

def init_state():
    if "base" not in st.session_state: st.session_state.base = load_base()
    if "supuestos" not in st.session_state: st.session_state.supuestos = load_csv(SUPUESTOS_PATH)
    if "rubrica" not in st.session_state: st.session_state.rubrica = load_rubrica()
    if "tipicidad" not in st.session_state: st.session_state.tipicidad = load_csv(TIPICIDAD_PATH)
    if "estandares" not in st.session_state: st.session_state.estandares = load_csv(ESTANDARES_PATH)
    if "juris" not in st.session_state: st.session_state.juris = load_csv(JURIS_PATH)
    if "catalogo_desc" not in st.session_state: st.session_state.catalogo_desc = load_csv(CATALOGO_DESC_PATH)
    if "hechos" not in st.session_state:
        st.session_state.hechos = st.session_state.supuestos.iloc[0].to_dict() if not st.session_state.supuestos.empty else {}
    if "auditoria" not in st.session_state: st.session_state.auditoria = []
    if "resultado" not in st.session_state: st.session_state.resultado = None
    if "sub" not in st.session_state: st.session_state.sub = None
    if "traz" not in st.session_state: st.session_state.traz = None
    if "prob" not in st.session_state: st.session_state.prob = pd.DataFrame()
    if "tipicidad_eval" not in st.session_state: st.session_state.tipicidad_eval = pd.DataFrame()
    if "conv_eval" not in st.session_state: st.session_state.conv_eval = pd.DataFrame()
    if "contradiccion" not in st.session_state: st.session_state.contradiccion = pd.DataFrame([
        {"Parte":"Fiscalía","Argumento":"","Evidencia":"","Impacto sugerido":""},
        {"Parte":"Defensa","Argumento":"","Evidencia":"","Impacto sugerido":""},
        {"Parte":"Comunidad indígena","Argumento":"","Evidencia":"","Impacto sugerido":""},
        {"Parte":"Tercero / empresa","Argumento":"","Evidencia":"","Impacto sugerido":""},
        {"Parte":"Amicus / Defensoría","Argumento":"","Evidencia":"","Impacto sugerido":""},
    ])

def log_event(action):
    st.session_state.auditoria.append({
        "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Acción": action,
        "Caso": st.session_state.hechos.get("nombre_caso",""),
        "Versión": "MCPI-IRCP-I v9"
    })

def fact_editor(hechos, prefix):
    c1,c2,c3 = st.columns(3)
    with c1:
        hechos["pais"] = st.text_input("País", hechos.get("pais",""), key=f"{prefix}_pais")
        hechos["norma"] = st.text_input("Norma", hechos.get("norma",""), key=f"{prefix}_norma")
        hechos["articulo"] = st.text_input("Artículo", hechos.get("articulo",""), key=f"{prefix}_art")
    with c2:
        hechos["tipo_penal"] = st.text_input("Tipo penal", hechos.get("tipo_penal",""), key=f"{prefix}_tipo")
        hechos["sujeto_colectivo"] = st.text_input("Sujeto colectivo", hechos.get("sujeto_colectivo",""), key=f"{prefix}_sujeto")
        hechos["tipo_protesta"] = st.text_input("Tipo de protesta", hechos.get("tipo_protesta",""), key=f"{prefix}_protesta")
    with c3:
        hechos["duracion"] = st.selectbox("Duración", ["Minutos","Varias horas","Más de un día"], index=["Minutos","Varias horas","Más de un día"].index(hechos.get("duracion","Varias horas")) if hechos.get("duracion","Varias horas") in ["Minutos","Varias horas","Más de un día"] else 1, key=f"{prefix}_dur")
        hechos["nivel_afectacion"] = st.selectbox("Nivel de afectación", ["Sin afectación","Interrupción parcial de circulación","Interrupción total","Afectación grave de servicios esenciales"], index=1, key=f"{prefix}_afec")
        hechos["contexto"] = st.text_input("Contexto", hechos.get("contexto",""), key=f"{prefix}_contexto")
    hechos["finalidad"] = st.text_area("Finalidad", hechos.get("finalidad",""), height=70, key=f"{prefix}_fin")
    hechos["descripcion"] = st.text_area("Descripción", hechos.get("descripcion",""), height=140, key=f"{prefix}_desc")
    opciones = ["Sí","No","No consta","No aplica","No / deficiente"]
    fields = [
        ("violencia_grave","Violencia grave"),("danos_personas","Daños a personas"),
        ("destruccion_bienes","Destrucción de bienes"),("uso_armas","Uso de armas"),
        ("organizacion_criminal","Organización criminal"),("riesgo_concreto_vida_integridad","Riesgo concreto vida/integridad"),
        ("dialogo_previo","Diálogo previo"),("consulta_previa","Consulta previa"),
        ("alternativas_no_penales","Alternativas no penales"),("prueba_individualizada","Prueba individualizada")
    ]
    for i in range(0,len(fields),2):
        c1,c2=st.columns(2)
        for col,(f,label) in zip([c1,c2], fields[i:i+2]):
            with col:
                val=hechos.get(f,"No")
                idx=opciones.index(val) if val in opciones else 1
                hechos[f]=st.selectbox(label, opciones, index=idx, key=f"{prefix}_{f}")
    return hechos

def render_result(res, hechos):
    st.subheader("Resultado MCPI-IRCP-I")
    st.dataframe(componentes(res), hide_index=True, use_container_width=True)
    c1,c2,c3=st.columns(3)
    c1.metric("CPI", res["CPI"])
    c2.metric("IRCP-I", res["IRCP-I"])
    c3.metric("Nivel / semáforo", f"{res['Nivel']} / {res['Semáforo']}")
    st.code(
        f"Formal = ({res['Legalidad']} × 0,60) + ({res['Claridad normativa']} × 0,40) = {res['Formal']}\n"
        f"Material = ({res['Lesividad']} × 0,60) + ({res['Mínima intervención penal']} × 0,40) = {res['Material']}\n"
        f"Europeo = ({res['Idoneidad']} × 0,50) + ({res['Necesidad penal']} × 0,50) = {res['Europeo']}\n"
        f"Alexy = [{res['Peso interés estatal']} ÷ ({res['Peso derechos indígenas']} + {res['Peso interés estatal']})] × 100 = {res['Alexy']}\n"
        f"CPI = {res['CPI']}\nIRCP-I = 100 − {res['CPI']} = {res['IRCP-I']}",
        language="text"
    )
    for r in reglas_cierre(res): st.warning(r)
    st.write(conclusion(hechos,res))
    fig, ax = plt.subplots(figsize=(8,4))
    vals = [res["Formal"],res["Material"],res["Europeo"],res["Alexy"],res["CPI"],res["IRCP-I"]]
    ax.bar(["Formal","Material","Europeo","Alexy","CPI","IRCP-I"], vals)
    ax.set_ylim(0,100); ax.grid(axis="y", linestyle="--", alpha=.4); st.pyplot(fig)

init_state()

st.title("MCPI–IRCP-I v9 · Motor judicial ampliado")
st.caption("Herramienta de apoyo: no reemplaza valoración judicial, prueba, contradicción ni motivación individual del caso.")

menu = st.sidebar.radio("Módulos", [
    "1. Expediente del caso",
    "2. Evaluación desde hechos",
    "3. Control probatorio",
    "4. Tipicidad estricta",
    "5. Control de convencionalidad",
    "6. Contradicción procesal",
    "7. Motivación judicial asistida",
    "8. Base jurisprudencial",
    "9. Comparador Argentina-Ecuador",
    "10. Dashboard de riesgos",
    "11. Análisis de sensibilidad",
    "12. Auditoría y trazabilidad",
    "13. Matriz de pesos editables",
    "14. Ideas para operadores de justicia",
    "15. Administración de bases",
])

if menu == "1. Expediente del caso":
    st.header("1. Expediente del caso")
    if not st.session_state.supuestos.empty:
        opts=[f"{i} | {r['nombre_caso']}" for i,r in st.session_state.supuestos.iterrows()]
        sel=st.selectbox("Cargar supuesto precargado", opts)
        idx=int(sel.split("|")[0].strip())
        if st.button("Cargar supuesto"):
            st.session_state.hechos=st.session_state.supuestos.loc[idx].to_dict()
            log_event("Cargar supuesto")
            st.success("Supuesto cargado.")
    st.subheader("Catálogo exhaustivo para el campo Descripción")
    st.info("Seleccione un tipo penal y cargue en el campo Descripción una recopilación amplia de supuestos comunes de criminalización de protesta indígena.")
    cat = st.session_state.catalogo_desc.copy() if "catalogo_desc" in st.session_state else pd.DataFrame()
    if not cat.empty:
        opts_cat = [f"{i} | {r['pais']} | {r['articulo']} | {r['tipo_penal']}" for i, r in cat.iterrows()]
        current_tipo = str(st.session_state.hechos.get("tipo_penal", "")).lower()
        default_cat_idx = 0
        for ii, opt in enumerate(opts_cat):
            if current_tipo and current_tipo in opt.lower():
                default_cat_idx = ii
                break
        sel_cat = st.selectbox("Catálogo de supuestos por tipo penal", opts_cat, index=default_cat_idx)
        idx_cat = int(sel_cat.split("|")[0].strip())
        with st.expander("Ver catálogo antes de cargar", expanded=False):
            st.text_area("Catálogo disponible", cat.loc[idx_cat, "catalogo_descripcion"], height=360, key="catalogo_preview")
        if st.button("Cargar catálogo exhaustivo en Descripción"):
            st.session_state.hechos["pais"] = cat.loc[idx_cat, "pais"]
            st.session_state.hechos["norma"] = cat.loc[idx_cat, "norma"]
            st.session_state.hechos["articulo"] = cat.loc[idx_cat, "articulo"]
            st.session_state.hechos["tipo_penal"] = cat.loc[idx_cat, "tipo_penal"]
            st.session_state.hechos["descripcion"] = cat.loc[idx_cat, "catalogo_descripcion"]
            log_event("Cargar catálogo exhaustivo en descripción")
            st.success("Catálogo cargado en el campo Descripción. Revise y edite el texto si corresponde.")

    st.session_state.hechos = fact_editor(dict(st.session_state.hechos), "exp")
    if st.button("Guardar expediente en sesión"):
        log_event("Guardar expediente")
        st.success("Expediente guardado.")

elif menu == "2. Evaluación desde hechos":
    st.header("2. Evaluación desde hechos")
    sub, traz, valores = score_desde_hechos(st.session_state.hechos, st.session_state.base)
    st.subheader("Matriz de trazabilidad hecho → variable")
    st.dataframe(traz, hide_index=True, use_container_width=True)
    st.subheader("Subcriterios sugeridos desde hechos")
    st.info("Puede editar tanto el valor como el peso de cada subcriterio. Para cada variable, la suma de pesos debe ser 1.00. Si no suma 1.00, la app lo advertirá y podrá normalizar automáticamente.")
    sub_edit=st.data_editor(sub, hide_index=True, use_container_width=True,
        column_config={"Valor": st.column_config.NumberColumn("Valor 0-100", min_value=0, max_value=100, step=1),
                       "Peso": st.column_config.NumberColumn("Peso editable", min_value=0.0, max_value=1.0, step=0.01, format="%.2f"),
                       "Aporte": st.column_config.NumberColumn("Aporte", disabled=True, format="%.2f"),
                       "Explicación metodológica": st.column_config.TextColumn(width="large")},
        disabled=["Modulo","Variable","Codigo","Subcriterio","Aporte","Explicación metodológica"])
    sub_edit["Valor"]=pd.to_numeric(sub_edit["Valor"], errors="coerce").fillna(0).clip(0,100)
    sub_edit["Peso"]=pd.to_numeric(sub_edit["Peso"], errors="coerce").fillna(0).clip(0,1)

    st.subheader("Control de pesos por variable")
    resumen_pesos = validar_pesos_subcriterios(sub_edit)
    st.dataframe(resumen_pesos, hide_index=True, use_container_width=True)

    colp1, colp2, colp3 = st.columns(3)
    with colp1:
        normalizar = st.button("Normalizar pesos automáticamente")
    with colp2:
        guardar_pesos = st.button("Guardar pesos editados en sesión")
    with colp3:
        restaurar = st.button("Restaurar pesos base del caso")

    if normalizar:
        sub_edit = normalizar_pesos_subcriterios(sub_edit)
        st.success("Pesos normalizados. La suma de pesos por variable ahora es 1.00.")
        log_event("Normalizar pesos de subcriterios")

    if guardar_pesos:
        st.session_state.sub = sub_edit.copy()
        st.success("Pesos editados guardados en la sesión.")
        log_event("Guardar pesos editados")

    if restaurar:
        sub_edit = sub.copy()
        st.success("Pesos restaurados al esquema base del modelo.")
        log_event("Restaurar pesos base")

    sub_edit, var_df, valores = calcular_variables_desde_subcriterios(sub_edit)

    if (validar_pesos_subcriterios(sub_edit)["Estado"] == "Revisar").any():
        st.warning("Hay variables cuyos pesos no suman 1.00. El cálculo se realiza igual, pero debe interpretarse como escenario metodológico alternativo.")

    st.subheader("Variables calculadas")
    st.dataframe(var_df, hide_index=True, use_container_width=True)
    res=calcular_indice(valores)
    st.session_state.resultado=res; st.session_state.sub=sub_edit; st.session_state.traz=traz
    log_event("Ejecutar evaluación desde hechos con pesos editables")
    render_result(res, st.session_state.hechos)

    st.download_button(
        "Descargar subcriterios con pesos editados CSV",
        sub_edit.to_csv(index=False).encode("utf-8-sig"),
        "subcriterios_pesos_editados.csv",
        "text/csv"
    )

elif menu == "3. Control probatorio":
    st.header("3. Control probatorio")
    default = pd.DataFrame([
        {"Hecho crítico":"Violencia grave","Estado":st.session_state.hechos.get("violencia_grave","No"),"Fuente probatoria":"","Impacto":"Lesividad / idoneidad"},
        {"Hecho crítico":"Daños a personas","Estado":st.session_state.hechos.get("danos_personas","No"),"Fuente probatoria":"","Impacto":"Lesividad"},
        {"Hecho crítico":"Destrucción de bienes","Estado":st.session_state.hechos.get("destruccion_bienes","No"),"Fuente probatoria":"","Impacto":"Lesividad"},
        {"Hecho crítico":"Uso de armas","Estado":st.session_state.hechos.get("uso_armas","No"),"Fuente probatoria":"","Impacto":"Riesgo concreto"},
        {"Hecho crítico":"Organización criminal","Estado":st.session_state.hechos.get("organizacion_criminal","No"),"Fuente probatoria":"","Impacto":"Tipicidad / criminalidad organizada"},
        {"Hecho crítico":"Riesgo concreto vida/integridad","Estado":st.session_state.hechos.get("riesgo_concreto_vida_integridad","No"),"Fuente probatoria":"","Impacto":"Proporcionalidad"},
        {"Hecho crítico":"Consulta previa","Estado":st.session_state.hechos.get("consulta_previa","No / deficiente"),"Fuente probatoria":"","Impacto":"Interculturalidad / mínima intervención"},
        {"Hecho crítico":"Autoría individualizada","Estado":st.session_state.hechos.get("prueba_individualizada","No consta"),"Fuente probatoria":"","Impacto":"Prueba / tipicidad"},
    ])
    edited=st.data_editor(default, num_rows="dynamic", hide_index=True, use_container_width=True,
        column_config={"Estado": st.column_config.SelectboxColumn(options=["Sí","No","No consta","No aplica","No / deficiente","Parcial","Alegado por Fiscalía","Alegado por defensa","Controvertido"])})
    st.session_state.prob=edited
    faltan=edited[(edited["Estado"].isin(["Sí","No / deficiente"])) & (edited["Fuente probatoria"].astype(str).str.len()==0)]
    if not faltan.empty:
        st.warning("Hay hechos relevantes sin fuente probatoria. Esto debe ser motivado antes de usar el resultado.")
    else:
        st.success("Control probatorio completo.")
    log_event("Actualizar control probatorio")

elif menu == "4. Tipicidad estricta":
    st.header("4. Control de tipicidad estricta")
    tip=st.session_state.tipicidad.copy()
    pais=st.session_state.hechos.get("pais","")
    tipo=st.session_state.hechos.get("tipo_penal","")
    filt=tip[(tip["pais"].astype(str).str.lower()==str(pais).lower()) & (tip["tipo_penal"].astype(str).str.lower()==str(tipo).lower())]
    if filt.empty:
        filt=tip[tip["pais"].astype(str).str.lower()==str(pais).lower()].head(8)
        st.info("No hay ficha exacta para este tipo penal; se muestran preguntas generales del país.")
    eval_df=filt[["elemento","pregunta_judicial","dimension"]].copy()
    eval_df["Respuesta"]="No consta"
    eval_df["Fuente / motivación"]=""
    edited=st.data_editor(eval_df, hide_index=True, use_container_width=True,
        column_config={"Respuesta": st.column_config.SelectboxColumn(options=["Sí","Parcial","No","No consta"]),
                       "pregunta_judicial": st.column_config.TextColumn(width="large")})
    st.session_state.tipicidad_eval=edited
    if (edited["Respuesta"]=="No").any() or (edited["Respuesta"]=="No consta").any():
        st.warning("Hay elementos del tipo penal no acreditados o sin constancia. Activar control estricto antes de imputar o condenar.")
    log_event("Actualizar tipicidad estricta")

elif menu == "5. Control de convencionalidad":
    st.header("5. Control de convencionalidad")
    df=st.session_state.estandares.copy()
    df["Resultado"]="No consta"
    df["Motivación / evidencia"]=""
    edited=st.data_editor(df, hide_index=True, use_container_width=True,
        column_config={"Resultado": st.column_config.SelectboxColumn(options=["Supera","No supera","Condicionado","No consta"]),
                       "pregunta_control": st.column_config.TextColumn(width="large"),
                       "impacto": st.column_config.TextColumn(width="large")})
    st.session_state.conv_eval=edited
    if (edited["Resultado"].isin(["No supera","No consta"])).any():
        st.warning("Existen estándares no superados o no motivados. Se requiere motivación reforzada.")
    log_event("Actualizar control de convencionalidad")

elif menu == "6. Contradicción procesal":
    st.header("6. Contradicción procesal")
    st.info("Registre posiciones de fiscalía, defensa, comunidad, terceros y amicus. La app no decide; estructura puntos controvertidos.")
    st.session_state.contradiccion=st.data_editor(st.session_state.contradiccion, num_rows="dynamic", hide_index=True, use_container_width=True)
    log_event("Actualizar contradicción procesal")

elif menu == "7. Motivación judicial asistida":
    st.header("7. Motivación judicial asistida")
    if st.session_state.resultado is None:
        sub,traz,valores=score_desde_hechos(st.session_state.hechos, st.session_state.base)
        res=calcular_indice(valores)
        st.session_state.sub=sub; st.session_state.traz=traz; st.session_state.resultado=res
    res=st.session_state.resultado
    st.subheader("Texto sugerido")
    texto=conclusion(st.session_state.hechos, res)
    st.write(texto)
    for r in reglas_cierre(res):
        st.warning(r)
    xls=to_excel({
        "Expediente": pd.DataFrame([st.session_state.hechos]),
        "Subcriterios": st.session_state.sub if st.session_state.sub is not None else pd.DataFrame(),
        "Trazabilidad": st.session_state.traz if st.session_state.traz is not None else pd.DataFrame(),
        "Probatorio": st.session_state.prob,
        "Tipicidad": st.session_state.tipicidad_eval,
        "Convencionalidad": st.session_state.conv_eval,
        "Contradiccion": st.session_state.contradiccion,
        "Resultado": pd.DataFrame([res]),
        "Componentes": componentes(res),
        "Reglas": pd.DataFrame({"Regla": reglas_cierre(res)}),
        "Auditoria": pd.DataFrame(st.session_state.auditoria),
    })
    st.download_button("Descargar expediente Excel completo", xls, "expediente_MCPI_IRCP_v8.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    doc=docx_report(st.session_state.hechos, res, st.session_state.sub, st.session_state.traz, st.session_state.prob, st.session_state.tipicidad_eval, st.session_state.conv_eval, st.session_state.contradiccion)
    if doc:
        st.download_button("Descargar motivación Word", doc, "motivacion_judicial_MCPI_IRCP_v8.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    log_event("Generar motivación judicial")

elif menu == "8. Base jurisprudencial":
    st.header("8. Base jurisprudencial y estándares")
    st.info("Base ampliada con casos de la tesis, estándares del Sistema Interamericano, precedentes de consulta previa, protesta, criminalización y control de convencionalidad. Las filas pueden editarse desde Administración de bases.")
    query=st.text_input("Buscar por tema, fuente, regla, tipo penal o estándar", "")
    df=st.session_state.juris.copy()
    if query:
        mask=df.astype(str).apply(lambda col: col.str.contains(query, case=False, na=False)).any(axis=1)
        df=df[mask]
    colj1, colj2, colj3 = st.columns(3)
    with colj1:
        fuente_sel = st.selectbox("Filtrar por fuente", ["Todas"] + sorted(st.session_state.juris["fuente"].astype(str).unique().tolist()) if "fuente" in st.session_state.juris.columns else ["Todas"])
    with colj2:
        impacto_sel = st.selectbox("Filtrar por impacto", ["Todos"] + sorted(st.session_state.juris["impacto_en_indice"].astype(str).unique().tolist()) if "impacto_en_indice" in st.session_state.juris.columns else ["Todos"])
    with colj3:
        st.metric("Registros visibles", len(df))
    if fuente_sel != "Todas" and "fuente" in df.columns:
        df = df[df["fuente"].astype(str) == fuente_sel]
    if impacto_sel != "Todos" and "impacto_en_indice" in df.columns:
        df = df[df["impacto_en_indice"].astype(str) == impacto_sel]
    st.dataframe(df, hide_index=True, use_container_width=True)
    st.download_button("Descargar base jurisprudencial filtrada CSV", df.to_csv(index=False).encode("utf-8-sig"), "base_jurisprudencial_estandares_filtrada.csv", "text/csv")
    st.warning("Nota metodológica: no todos los eventos de criminalización son sentencias. Cuando no exista sentencia identificada, deben clasificarse como casos documentados de judicialización o criminalización, no como jurisprudencia.")

elif menu == "9. Comparador Argentina-Ecuador":
    st.header("9. Comparador Argentina–Ecuador")
    base=st.session_state.base.copy()
    st.dataframe(base[["País","Artículo","Tipo penal","Categoría","Legalidad","Claridad normativa","Lesividad","Mínima intervención penal","Idoneidad","Necesidad penal"]], hide_index=True, use_container_width=True)
    prom=base.groupby("País", as_index=False)[NUMERIC_COLS].mean()
    st.subheader("Promedio por país")
    st.dataframe(prom, hide_index=True, use_container_width=True)
    fig, ax = plt.subplots(figsize=(7,4))
    vals=[]
    labels=[]
    for _,r in prom.iterrows():
        res=calcular_indice({c:r[c] for c in NUMERIC_COLS})
        vals.append(res["IRCP-I"]); labels.append(r["País"])
    ax.bar(labels, vals); ax.set_ylim(0,100); ax.set_ylabel("IRCP-I promedio"); st.pyplot(fig)

elif menu == "10. Dashboard de riesgos":
    st.header("10. Dashboard de riesgos")
    rows=[]
    for _,r in st.session_state.base.iterrows():
        res=calcular_indice({c:r[c] for c in NUMERIC_COLS})
        rows.append({"País":r["País"],"Artículo":r["Artículo"],"Tipo penal":r["Tipo penal"],"CPI":res["CPI"],"IRCP-I":res["IRCP-I"],"Nivel":res["Nivel"],"Semáforo":res["Semáforo"]})
    df=pd.DataFrame(rows).sort_values("IRCP-I", ascending=False)
    st.dataframe(df, hide_index=True, use_container_width=True)
    fig, ax = plt.subplots(figsize=(10, max(5, len(df)*.28)))
    ax.barh(df["País"]+" - "+df["Tipo penal"], df["IRCP-I"])
    ax.set_xlim(0,100); ax.set_xlabel("IRCP-I"); ax.grid(axis="x", linestyle="--", alpha=.4); st.pyplot(fig)

elif menu == "11. Análisis de sensibilidad":
    st.header("11. Análisis de sensibilidad")
    if st.session_state.resultado is None:
        sub,traz,valores=score_desde_hechos(st.session_state.hechos, st.session_state.base)
        res=calcular_indice(valores)
    else:
        res=st.session_state.resultado
    st.write("Modifique pesos. Deben sumar 100%.")
    c1,c2,c3,c4=st.columns(4)
    wF=c1.number_input("Formal", 0.0, 1.0, 0.25, step=0.05)
    wM=c2.number_input("Material", 0.0, 1.0, 0.30, step=0.05)
    wE=c3.number_input("Europeo", 0.0, 1.0, 0.25, step=0.05)
    wA=c4.number_input("Alexy", 0.0, 1.0, 0.20, step=0.05)
    total=wF+wM+wE+wA
    if abs(total-1)>0.001: st.error(f"Los pesos suman {total:.2f}. Deben sumar 1.00.")
    cpi=round(res["Formal"]*wF+res["Material"]*wM+res["Europeo"]*wE+res["Alexy"]*wA,2)
    ircp=round(100-cpi,2); nivel,sem=nivel_riesgo(ircp)
    st.metric("IRCP-I recalculado", ircp, f"{nivel} / {sem}")

elif menu == "12. Auditoría y trazabilidad":
    st.header("12. Auditoría y trazabilidad")
    st.dataframe(pd.DataFrame(st.session_state.auditoria), hide_index=True, use_container_width=True)
    st.download_button("Descargar auditoría CSV", pd.DataFrame(st.session_state.auditoria).to_csv(index=False).encode("utf-8-sig"), "auditoria_MCPI_IRCP.csv", "text/csv")

elif menu == "13. Matriz de pesos editables":
    st.header("13. Matriz de pesos editables")
    st.info("Este módulo permite revisar, editar, validar y exportar los pesos de subcriterios usados en el caso activo.")
    if st.session_state.sub is None or st.session_state.sub.empty:
        sub, traz, valores = score_desde_hechos(st.session_state.hechos, st.session_state.base)
        st.session_state.sub = sub
        st.session_state.traz = traz

    editable = st.data_editor(
        st.session_state.sub,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={
            "Peso": st.column_config.NumberColumn("Peso editable", min_value=0.0, max_value=1.0, step=0.01, format="%.2f"),
            "Valor": st.column_config.NumberColumn("Valor 0-100", min_value=0, max_value=100, step=1),
            "Aporte": st.column_config.NumberColumn("Aporte", disabled=True, format="%.2f"),
            "Explicación metodológica": st.column_config.TextColumn(width="large"),
        },
        disabled=["Aporte"]
    )
    editable["Peso"] = pd.to_numeric(editable["Peso"], errors="coerce").fillna(0).clip(0, 1)
    editable["Valor"] = pd.to_numeric(editable["Valor"], errors="coerce").fillna(0).clip(0, 100)
    editable["Aporte"] = (editable["Peso"] * editable["Valor"]).round(2)

    st.subheader("Validación de suma de pesos")
    resumen = validar_pesos_subcriterios(editable)
    st.dataframe(resumen, hide_index=True, use_container_width=True)

    c1, c2, c3 = st.columns(3)
    if c1.button("Aplicar pesos a la sesión"):
        st.session_state.sub = editable.copy()
        _, var_df, valores = calcular_variables_desde_subcriterios(editable)
        st.session_state.resultado = calcular_indice(valores)
        st.success("Pesos aplicados a la sesión.")
        log_event("Aplicar matriz de pesos editable")
    if c2.button("Normalizar y aplicar"):
        editable = normalizar_pesos_subcriterios(editable)
        st.session_state.sub = editable.copy()
        _, var_df, valores = calcular_variables_desde_subcriterios(editable)
        st.session_state.resultado = calcular_indice(valores)
        st.success("Pesos normalizados y aplicados.")
        log_event("Normalizar y aplicar matriz de pesos")
    if c3.button("Recalcular resultado"):
        _, var_df, valores = calcular_variables_desde_subcriterios(editable)
        res = calcular_indice(valores)
        st.session_state.resultado = res
        render_result(res, st.session_state.hechos)

    st.download_button(
        "Descargar matriz de pesos CSV",
        editable.to_csv(index=False).encode("utf-8-sig"),
        "matriz_pesos_subcriterios_MCPI_IRCP.csv",
        "text/csv"
    )
    st.download_button(
        "Descargar matriz de pesos Excel",
        to_excel({"Pesos_subcriterios": editable, "Validacion_pesos": resumen}),
        "matriz_pesos_subcriterios_MCPI_IRCP.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

elif menu == "14. Ideas para operadores de justicia":
    st.header("14. Ideas para operadores de justicia")
    ideas = pd.DataFrame([
        {"Mejora": "Modo juez", "Descripción": "Formulario breve para resolver caso concreto: hechos, prueba, tipicidad, convencionalidad y motivación."},
        {"Mejora": "Modo fiscalía", "Descripción": "Filtro previo de imputación: exige daño grave, prueba individualizada, dolo específico y alternativas no penales agotadas."},
        {"Mejora": "Modo defensa / amicus", "Descripción": "Genera argumentos sobre mínima intervención, protesta protegida, consulta previa y enfoque intercultural."},
        {"Mejora": "Matriz de prueba", "Descripción": "Cada hecho crítico debe estar vinculado a fuente probatoria, grado de certeza y contradicción procesal."},
        {"Mejora": "Alertas de criminalización", "Descripción": "Alertas automáticas cuando se use terrorismo, sabotaje, asociación ilícita o rebelión sin violencia grave."},
        {"Mejora": "Control de consulta previa", "Descripción": "Checklist específico para verificar si el conflicto surge de falta de consulta o afectación territorial."},
        {"Mejora": "Control de medidas cautelares", "Descripción": "Módulo para evaluar prisión preventiva, detención, estado de excepción y uso de fuerza."},
        {"Mejora": "Repositorio de precedentes", "Descripción": "Búsqueda por país, delito, derecho afectado, estándar SIDH y utilidad para decisión judicial."},
        {"Mejora": "Informe judicial con doble versión", "Descripción": "Versión corta para audiencia y versión larga para sentencia o dictamen."},
        {"Mejora": "Auditoría institucional", "Descripción": "Registro de quién cambió pesos, hechos, fuentes y conclusiones para trazabilidad."},
        {"Mejora": "Validación experta", "Descripción": "Encuesta para jueces, fiscales, defensores, académicos y autoridades indígenas sobre pesos y resultados."},
        {"Mejora": "Panel de riesgos por territorio", "Descripción": "Mapa o tablero por provincia, tipo penal y contexto extractivo/territorial."},
    ])
    st.dataframe(ideas, hide_index=True, use_container_width=True)
    st.download_button("Descargar ideas CSV", ideas.to_csv(index=False).encode("utf-8-sig"), "ideas_mejora_operadores_justicia.csv", "text/csv")
    st.markdown("""
### Recomendación de uso institucional

La app debe presentarse como herramienta de apoyo y no como sistema de decisión automática. Para uso por operadores de justicia, cada resultado debe estar respaldado por:

1. hecho acreditado;
2. fuente probatoria;
3. elemento típico;
4. estándar constitucional o interamericano;
5. subcriterio afectado;
6. fórmula aplicada;
7. motivación humana final.
""")

else:
    st.header("15. Administración de bases")
    tabs=st.tabs(["Delitos","Supuestos","Subcriterios","Tipicidad","Convencionalidad","Jurisprudencia","Catálogo descripciones"])
    with tabs[0]:
        st.session_state.base=st.data_editor(st.session_state.base, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar delitos CSV", st.session_state.base.to_csv(index=False).encode("utf-8-sig"), "tipos_penales_base_ampliada.csv", "text/csv")
    with tabs[1]:
        st.session_state.supuestos=st.data_editor(st.session_state.supuestos, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar supuestos CSV", st.session_state.supuestos.to_csv(index=False).encode("utf-8-sig"), "supuestos_hecho_base.csv", "text/csv")
    with tabs[2]:
        st.dataframe(st.session_state.rubrica, hide_index=True, use_container_width=True)
    with tabs[3]:
        st.session_state.tipicidad=st.data_editor(st.session_state.tipicidad, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar tipicidad CSV", st.session_state.tipicidad.to_csv(index=False).encode("utf-8-sig"), "tipicidad_por_delito.csv", "text/csv")
    with tabs[4]:
        st.session_state.estandares=st.data_editor(st.session_state.estandares, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar convencionalidad CSV", st.session_state.estandares.to_csv(index=False).encode("utf-8-sig"), "estandares_convencionalidad.csv", "text/csv")
    with tabs[5]:
        st.session_state.juris=st.data_editor(st.session_state.juris, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar jurisprudencia CSV", st.session_state.juris.to_csv(index=False).encode("utf-8-sig"), "jurisprudencia_base.csv", "text/csv")
    with tabs[6]:
        if "catalogo_desc" not in st.session_state:
            st.session_state.catalogo_desc = load_csv(CATALOGO_DESC_PATH)
        st.session_state.catalogo_desc=st.data_editor(st.session_state.catalogo_desc, num_rows="dynamic", hide_index=True, use_container_width=True)
        st.download_button("Descargar catálogo de descripciones CSV", st.session_state.catalogo_desc.to_csv(index=False).encode("utf-8-sig"), "catalogo_descripciones_tipo_penal.csv", "text/csv")

st.caption("MCPI–IRCP-I v9. Herramienta académica y judicial de apoyo con pesos editables, supuestos ampliados y base jurisprudencial reforzada. La decisión final corresponde exclusivamente a la autoridad competente.")
