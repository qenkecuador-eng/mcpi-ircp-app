from pathlib import Path
from io import BytesIO
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt

try:
    from docx import Document
except Exception:
    Document = None

APP_DIR = Path(__file__).parent
BASE_PATH = APP_DIR / "tipos_penales_base_ampliada.csv"
SUBCRITERIOS_PATH = APP_DIR / "subcriterios_MCPI_IRCP_app_unico.xlsx"
SUPUESTOS_PATH = APP_DIR / "supuestos_hecho_base.csv"

st.set_page_config(
    page_title="Matriz MCPI-IRCP-I v7",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

COLUMNAS_BASE = [
    "País", "Norma", "Artículo", "Tipo penal", "Categoría", "Estado",
    "Contexto típico de protesta", "Legalidad", "Claridad normativa",
    "Lesividad", "Mínima intervención penal", "Idoneidad", "Necesidad penal",
    "Peso derechos indígenas", "Peso interés estatal", "Observación metodológica"
]

NUMERIC_COLS = [
    "Legalidad", "Claridad normativa", "Lesividad", "Mínima intervención penal",
    "Idoneidad", "Necesidad penal", "Peso derechos indígenas", "Peso interés estatal"
]

RUBRICA_COLS = [
    "Modulo", "Variable", "Codigo", "Subcriterio", "Peso_sobre_variable",
    "Valor_editable_0_100", "Aporte_calculado", "Formula_de_aporte",
    "Formula_de_variable", "Explicacion_metodologica", "Puntaje_alto_75_100",
    "Puntaje_medio_40_74", "Puntaje_bajo_0_39", "Fuente_metodologica_base"
]

VAR_CANONICA = {
    "Legalidad / tipicidad estricta": "Legalidad",
    "Legalidad": "Legalidad",
    "Claridad normativa": "Claridad normativa",
    "Lesividad": "Lesividad",
    "Mínima intervención penal": "Mínima intervención penal",
    "Minima intervencion penal": "Mínima intervención penal",
    "Idoneidad": "Idoneidad",
    "Necesidad penal": "Necesidad penal",
    "Peso derechos indígenas": "Peso derechos indígenas",
    "Peso derechos indigenas": "Peso derechos indígenas",
    "Peso interés estatal": "Peso interés estatal",
    "Peso interes estatal": "Peso interés estatal",
}

def normalizar_base(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for c in COLUMNAS_BASE:
        if c not in df.columns:
            df[c] = 0 if c in NUMERIC_COLS else ""
    df = df[COLUMNAS_BASE]
    for c in NUMERIC_COLS:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).clip(0, 100)
    for c in [x for x in COLUMNAS_BASE if x not in NUMERIC_COLS]:
        df[c] = df[c].fillna("").astype(str)
    return df

def normalizar_rubrica(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    for c in RUBRICA_COLS:
        if c not in df.columns:
            df[c] = 0 if c in ["Peso_sobre_variable", "Valor_editable_0_100", "Aporte_calculado"] else ""
    df = df[RUBRICA_COLS]
    df["Peso_sobre_variable"] = pd.to_numeric(df["Peso_sobre_variable"], errors="coerce").fillna(0)
    df["Valor_editable_0_100"] = pd.to_numeric(df["Valor_editable_0_100"], errors="coerce").fillna(0).clip(0, 100)
    df["Aporte_calculado"] = (df["Valor_editable_0_100"] * df["Peso_sobre_variable"]).round(2)
    for c in [x for x in RUBRICA_COLS if x not in ["Peso_sobre_variable", "Valor_editable_0_100", "Aporte_calculado"]]:
        df[c] = df[c].fillna("").astype(str)
    return df

def cargar_base_default() -> pd.DataFrame:
    if BASE_PATH.exists():
        return normalizar_base(pd.read_csv(BASE_PATH))
    return normalizar_base(pd.DataFrame(columns=COLUMNAS_BASE))

def cargar_rubrica_default() -> pd.DataFrame:
    if SUBCRITERIOS_PATH.exists():
        return normalizar_rubrica(pd.read_excel(SUBCRITERIOS_PATH, sheet_name=0))
    return normalizar_rubrica(pd.DataFrame(columns=RUBRICA_COLS))

def cargar_supuestos_default() -> pd.DataFrame:
    if SUPUESTOS_PATH.exists():
        return pd.read_csv(SUPUESTOS_PATH).fillna("")
    return pd.DataFrame()

def nivel_compatibilidad(valor):
    valor = float(valor)
    if valor >= 75:
        return "Compatibilidad alta", "Verde"
    if valor >= 60:
        return "Compatibilidad condicionada", "Amarillo"
    if valor >= 40:
        return "Compatibilidad débil", "Naranja"
    return "Incompatibilidad", "Rojo"

def nivel_riesgo(valor):
    valor = float(valor)
    if valor >= 70:
        return "Riesgo alto", "Rojo"
    if valor >= 50:
        return "Riesgo medio-alto", "Naranja"
    if valor >= 35:
        return "Riesgo medio", "Amarillo"
    return "Riesgo bajo", "Verde"

def calcular_indice_desde_variables(valores: dict) -> dict:
    legalidad = float(valores.get("Legalidad", 0))
    claridad = float(valores.get("Claridad normativa", 0))
    lesividad = float(valores.get("Lesividad", 0))
    minima = float(valores.get("Mínima intervención penal", 0))
    idoneidad = float(valores.get("Idoneidad", 0))
    necesidad = float(valores.get("Necesidad penal", 0))
    pdi = float(valores.get("Peso derechos indígenas", 0))
    pie = float(valores.get("Peso interés estatal", 0))

    formal = round(legalidad * 0.60 + claridad * 0.40, 2)
    material = round(lesividad * 0.60 + minima * 0.40, 2)
    europeo = round(idoneidad * 0.50 + necesidad * 0.50, 2)
    alexy = round((pie / (pdi + pie)) * 100, 2) if (pdi + pie) else 0
    peso_relativo = round(pdi / pie, 2) if pie else 999
    cpi = round((formal * 0.25) + (material * 0.30) + (europeo * 0.25) + (alexy * 0.20), 2)
    ircp = round(100 - cpi, 2)
    nivel, sem = nivel_riesgo(ircp)

    return {
        "Legalidad": legalidad,
        "Claridad normativa": claridad,
        "Lesividad": lesividad,
        "Mínima intervención penal": minima,
        "Idoneidad": idoneidad,
        "Necesidad penal": necesidad,
        "Peso derechos indígenas": pdi,
        "Peso interés estatal": pie,
        "Formal": formal,
        "Material": material,
        "Europeo": europeo,
        "Alexy": alexy,
        "Peso relativo": peso_relativo,
        "CPI": cpi,
        "IRCP-I": ircp,
        "Nivel": nivel,
        "Semáforo": sem,
    }

def calcular_fila(row) -> dict:
    valores = {c: row[c] for c in NUMERIC_COLS}
    res = calcular_indice_desde_variables(valores)
    return {
        "País": row["País"],
        "Artículo": row["Artículo"],
        "Tipo penal": row["Tipo penal"],
        "Categoría": row["Categoría"],
        "Estado": row["Estado"],
        "Formal": res["Formal"],
        "Material": res["Material"],
        "Europeo": res["Europeo"],
        "Peso relativo": res["Peso relativo"],
        "Alexy": res["Alexy"],
        "CPI": res["CPI"],
        "IRCP-I": res["IRCP-I"],
        "Nivel": res["Nivel"],
        "Semáforo": res["Semáforo"],
    }

def base_resultados(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    return pd.DataFrame([calcular_fila(row) for _, row in df.iterrows()])

def buscar_valores_formales(base: pd.DataFrame, pais: str, articulo: str, tipo_penal: str) -> dict:
    # Para el supuesto de art. 194 CP se toman valores de norma en abstracto.
    mask = (
        base["País"].str.lower().eq(str(pais).lower()) &
        (base["Artículo"].str.contains(str(articulo).replace("art. ", "").replace(" CP", ""), case=False, na=False) |
         base["Tipo penal"].str.lower().eq(str(tipo_penal).lower()))
    )
    if mask.any():
        row = base[mask].iloc[0]
        return {
            "Legalidad": float(row["Legalidad"]),
            "Claridad normativa": float(row["Claridad normativa"]),
        }
    return {"Legalidad": 55.0, "Claridad normativa": 50.0}

def si(valor) -> bool:
    return str(valor).strip().lower() in ["sí", "si", "yes", "true", "1"]

def no(valor) -> bool:
    return str(valor).strip().lower() in ["no", "false", "0"]

def scoring_desde_hechos(hechos: dict, base: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, dict]:
    """
    Convierte hechos en subcriterios sugeridos.
    La lógica está calibrada para reproducir el supuesto trabajado:
    comunidad indígena + proyecto extractivo inconsulto + corte parcial temporal sin violencia grave.
    """
    formales = buscar_valores_formales(
        base,
        hechos.get("pais", ""),
        hechos.get("articulo", ""),
        hechos.get("tipo_penal", ""),
    )

    violencia = si(hechos.get("violencia_grave", "No"))
    danos_personas = si(hechos.get("danos_personas", "No"))
    destruccion = si(hechos.get("destruccion_bienes", "No"))
    armas = si(hechos.get("uso_armas", "No"))
    org_criminal = si(hechos.get("organizacion_criminal", "No"))
    riesgo_vida = si(hechos.get("riesgo_concreto_vida_integridad", "No"))
    indigena = "indígena" in str(hechos.get("sujeto_colectivo", "")).lower() or "indigena" in str(hechos.get("sujeto_colectivo", "")).lower()
    extractivo = "extract" in str(hechos.get("contexto", "")).lower()
    consulta_def = "no" in str(hechos.get("consulta_previa", "")).lower() or "deficiente" in str(hechos.get("consulta_previa", "")).lower()
    afectacion_parcial = "parcial" in str(hechos.get("nivel_afectacion", "")).lower()
    varias_horas = "varias" in str(hechos.get("duracion", "")).lower() or "horas" in str(hechos.get("duracion", "")).lower()
    alternativas = si(hechos.get("alternativas_no_penales", "Sí"))

    # Valores sugeridos base, calibrados al supuesto doctoral.
    les_dano = 35 if afectacion_parcial else 50
    les_gravedad = 25
    les_riesgo = 20
    les_prueba = 35

    if violencia:
        les_gravedad += 25
    if danos_personas:
        les_dano += 25
        les_gravedad += 20
    if destruccion:
        les_dano += 15
        les_gravedad += 10
    if armas:
        les_gravedad += 15
        les_riesgo += 20
    if riesgo_vida:
        les_riesgo = 70
        les_gravedad += 20
    if org_criminal:
        les_prueba += 15

    les_dano = min(100, les_dano)
    les_gravedad = min(100, les_gravedad)
    les_riesgo = min(100, les_riesgo)
    les_prueba = min(100, les_prueba)

    min_dialogo = 20 if "no" in str(hechos.get("dialogo_previo", "")).lower() or "consta" in str(hechos.get("dialogo_previo", "")).lower() else 60
    min_consulta = 15 if consulta_def else 60
    min_medidas = 30 if alternativas else 65
    min_ultima = 20 if alternativas else 65

    ido_finalidad = 60
    ido_conexion = 35
    ido_riesgo = 20 if not riesgo_vida else 70
    ido_violencia = 10
    ido_delimitacion = 30
    ido_evidencia = 40
    ido_intercultural = 15 if indigena else 35

    if violencia or danos_personas or destruccion or armas:
        ido_violencia = 55
        ido_conexion += 15
    if riesgo_vida:
        ido_riesgo = 75
    if not indigena:
        ido_intercultural = 35

    nec_dialogo = 15 if min_dialogo <= 20 else 55
    nec_consulta = 10 if consulta_def and indigena else 50
    nec_mediacion = 20 if indigena else 35
    nec_admin = 30 if alternativas else 65
    nec_preventivas = 35 if alternativas else 60
    nec_eficacia = 35 if alternativas else 60
    nec_no_agotamiento = 15 if alternativas else 55

    pdi_territorial = 90 if indigena and extractivo else (70 if indigena else 35)
    pdi_abstracto = 90 if indigena else 50
    pdi_fiabilidad = 75 if indigena else 40

    pie_intensidad = 35 if afectacion_parcial else 55
    if violencia or danos_personas or destruccion or armas or riesgo_vida:
        pie_intensidad += 25
    pie_abstracto = 55
    pie_fiabilidad = 30
    if varias_horas:
        pie_fiabilidad = 30
    if riesgo_vida or danos_personas:
        pie_fiabilidad += 25

    rows = [
        ["Control formal", "Legalidad", "LEG-BASE", "Valor abstracto del tipo penal seleccionado", 1.00, formales["Legalidad"], "Se toma de la base de delitos porque el control formal evalúa la norma en abstracto."],
        ["Control formal", "Claridad normativa", "CLA-BASE", "Valor abstracto del tipo penal seleccionado", 1.00, formales["Claridad normativa"], "Se toma de la base de delitos porque la claridad normativa evalúa el texto del tipo penal."],

        ["Control material", "Lesividad", "LES-01", "Daño real", 0.30, les_dano, "La interrupción parcial de circulación genera afectación real, pero no daño penal grave."],
        ["Control material", "Lesividad", "LES-02", "Gravedad del daño", 0.25, les_gravedad, "No se acreditan violencia grave, daños personales, armas ni destrucción de bienes."],
        ["Control material", "Lesividad", "LES-03", "Concreción del riesgo", 0.25, les_riesgo, "No se acredita riesgo concreto para vida o integridad física."],
        ["Control material", "Lesividad", "LES-04", "Prueba individualizada", 0.20, les_prueba, "No consta individualización robusta de autoría penal ni daño atribuible individualmente."],

        ["Control material", "Mínima intervención penal", "MIN-01", "Agotamiento de diálogo", 0.25, min_dialogo, "La protesta busca diálogo; si no se demuestra diálogo previo, baja la compatibilidad penal."],
        ["Control material", "Mínima intervención penal", "MIN-02", "Consulta previa o vía institucional", 0.25, min_consulta, "El conflicto se origina en un proyecto extractivo considerado inconsulto."],
        ["Control material", "Mínima intervención penal", "MIN-03", "Existencia de medidas no penales", 0.25, min_medidas, "Existen alternativas como diálogo, mediación, consulta, gestión de tránsito o vías administrativas."],
        ["Control material", "Mínima intervención penal", "MIN-04", "Derecho penal como última ratio", 0.25, min_ultima, "No se evidencia que la vía penal fuera estrictamente indispensable."],

        ["Test europeo", "Idoneidad", "IDO-01", "Finalidad legítima", 0.15, ido_finalidad, "El Estado puede proteger circulación, orden público o servicios."],
        ["Test europeo", "Idoneidad", "IDO-02", "Conexión medio-fin", 0.15, ido_conexion, "La vía penal tiene conexión débil si el conflicto requiere diálogo y consulta."],
        ["Test europeo", "Idoneidad", "IDO-03", "Riesgo real actual individualizado", 0.20, ido_riesgo, "No hay riesgo concreto individualizado para vida o integridad."],
        ["Test europeo", "Idoneidad", "IDO-04", "Violencia grave o lesividad relevante", 0.20, ido_violencia, "No se acreditan violencia grave ni daños relevantes."],
        ["Test europeo", "Idoneidad", "IDO-05", "Delimitación estricta de aplicación", 0.10, ido_delimitacion, "Debe distinguirse protesta protegida de conducta penalmente punible."],
        ["Test europeo", "Idoneidad", "IDO-06", "Evidencia fáctica suficiente", 0.10, ido_evidencia, "Hay evidencia de interrupción, pero no de daño grave o riesgo concreto."],
        ["Test europeo", "Idoneidad", "IDO-07", "Enfoque intercultural inicial", 0.10, ido_intercultural, "La activación penal muestra baja consideración inicial del contexto indígena."],

        ["Test europeo", "Necesidad penal", "NEC-01", "Diálogo previo", 0.15, nec_dialogo, "La finalidad de la protesta es exigir diálogo."],
        ["Test europeo", "Necesidad penal", "NEC-02", "Consulta previa adecuada", 0.20, nec_consulta, "La consulta previa aparece inexistente o deficiente según el supuesto."],
        ["Test europeo", "Necesidad penal", "NEC-03", "Mediación intercultural", 0.15, nec_mediacion, "Debió evaluarse mediación intercultural antes de penalizar."],
        ["Test europeo", "Necesidad penal", "NEC-04", "Medidas administrativas o civiles", 0.15, nec_admin, "Había medidas no penales de gestión del conflicto."],
        ["Test europeo", "Necesidad penal", "NEC-05", "Medidas preventivas no penales", 0.10, nec_preventivas, "Podían adoptarse medidas de gestión preventiva de tránsito y diálogo."],
        ["Test europeo", "Necesidad penal", "NEC-06", "Eficacia de alternativas menos lesivas", 0.15, nec_eficacia, "Las alternativas podían proteger la circulación sin criminalizar."],
        ["Test europeo", "Necesidad penal", "NEC-07", "No agotamiento real de alternativas", 0.10, nec_no_agotamiento, "No se demuestra agotamiento real de alternativas no penales."],

        ["Alexy", "Peso derechos indígenas", "PDI-01", "Intensidad de afectación territorial/cultural", 0.35, pdi_territorial, "Proyecto extractivo, territorio y consulta previa elevan el peso del derecho indígena."],
        ["Alexy", "Peso derechos indígenas", "PDI-02", "Peso abstracto del derecho indígena comprometido", 0.35, pdi_abstracto, "Territorio, consulta, protesta, participación y autodeterminación tienen peso constitucional/convencional alto."],
        ["Alexy", "Peso derechos indígenas", "PDI-03", "Fiabilidad probatoria del contexto indígena", 0.30, pdi_fiabilidad, "El supuesto identifica comunidad indígena, proyecto extractivo y reclamo de consulta."],

        ["Alexy", "Peso interés estatal", "PIE-01", "Intensidad de afectación al interés estatal", 0.35, min(100, pie_intensidad), "Existe afectación parcial de circulación, pero sin daño grave ni riesgo concreto."],
        ["Alexy", "Peso interés estatal", "PIE-02", "Peso abstracto del interés estatal", 0.35, pie_abstracto, "La circulación y el orden público son fines legítimos, pero no absolutos."],
        ["Alexy", "Peso interés estatal", "PIE-03", "Fiabilidad probatoria del daño estatal", 0.30, min(100, pie_fiabilidad), "Se acredita interrupción, pero no daño grave, armas ni riesgo para terceros."],
    ]

    sub = pd.DataFrame(rows, columns=["Modulo", "Variable", "Codigo", "Subcriterio", "Peso", "Valor sugerido", "Explicación metodológica"])
    sub["Aporte"] = (sub["Peso"] * sub["Valor sugerido"]).round(2)

    variables = sub.groupby("Variable", as_index=False)["Aporte"].sum().rename(columns={"Aporte": "Valor"})
    valores = {r["Variable"]: float(r["Valor"]) for _, r in variables.iterrows()}

    trazabilidad_rows = [
        ["Corte temporal de ruta", "Acreditado", "Tipo penal / interés estatal", "Activa análisis del art. 194 CP y eleva moderadamente interés estatal."],
        ["Interrupción parcial durante varias horas", "Acreditado", "Lesividad / interés estatal", "Reconoce afectación real, pero no daño penal grave automático."],
        ["Comunidad indígena", "Acreditado", "Alexy / derechos indígenas", "Eleva peso de derechos colectivos, territorio, participación y consulta."],
        ["Proyecto extractivo considerado inconsulto", "Acreditado / alegado", "Mínima intervención / necesidad / Alexy", "Refuerza exigencia de consulta previa y vías no penales."],
        ["Finalidad de exigir diálogo y consulta previa", "Acreditado", "Necesidad penal / mínima intervención", "Reduce justificación de la vía penal como primera respuesta."],
        ["Violencia grave", "No acreditado", "Lesividad / idoneidad", "Reduce gravedad del daño y conexión penal medio-fin."],
        ["Daños a personas", "No acreditado", "Lesividad", "Reduce daño penal relevante."],
        ["Destrucción de bienes", "No acreditado", "Lesividad", "Reduce daño material penalmente relevante."],
        ["Uso de armas", "No acreditado", "Idoneidad / riesgo", "Reduce riesgo concreto y gravedad."],
        ["Organización criminal", "No acreditado", "Prueba individualizada / agravación", "Evita asimilar organización comunitaria con criminalidad."],
        ["Riesgo concreto para vida o integridad", "No acreditado", "Lesividad / proporcionalidad", "Reduce fuertemente la justificación penal."],
    ]
    traz = pd.DataFrame(trazabilidad_rows, columns=["Hecho del caso", "Estado", "Variable afectada", "Impacto metodológico"])
    return sub, traz, valores

def reglas_de_cierre(res: dict) -> list[str]:
    reglas = []
    if res["Legalidad"] < 40:
        reglas.append("Legalidad menor a 40: presunción de incompatibilidad formal.")
    if res["Lesividad"] < 40:
        reglas.append("Lesividad menor a 40: no se acredita daño penal grave suficiente.")
    if res["Mínima intervención penal"] < 40:
        reglas.append("Mínima intervención menor a 40: debieron priorizarse vías no penales.")
    if res["Idoneidad"] < 75:
        reglas.append("Idoneidad menor a 75: la criminalización penal no queda plenamente justificada.")
    if res["Necesidad penal"] < 60:
        reglas.append("Necesidad penal menor a 60: existen o debieron agotarse alternativas menos lesivas.")
    if res["Peso relativo"] > 1:
        reglas.append("Alexy favorece derechos indígenas/protesta: exige motivación reforzada.")
    if res["IRCP-I"] >= 70:
        reglas.append("IRCP-I alto: riesgo severo de criminalización penal intercultural.")
    elif res["IRCP-I"] >= 50:
        reglas.append("IRCP-I medio-alto: riesgo relevante de criminalización penal intercultural.")
    return reglas

def conclusion_automatica(hechos: dict, res: dict) -> str:
    return (
        f"En el supuesto analizado, la aplicación de {hechos.get('articulo','')} "
        f"({hechos.get('tipo_penal','')}) a {hechos.get('sujeto_colectivo','')} que realiza "
        f"{hechos.get('tipo_protesta','')} para {hechos.get('finalidad','')} arroja una "
        f"Compatibilidad Penal Intercultural de {res['CPI']}/100 y un IRCP-I de {res['IRCP-I']}/100, "
        f"equivalente a {res['Nivel']} / semáforo {res['Semáforo']}. El resultado se explica porque, "
        f"aunque existe afectación a la circulación o al interés estatal invocado, no se acreditan violencia "
        f"grave, daño a personas, destrucción de bienes, uso de armas, organización criminal ni riesgo concreto "
        f"para vida o integridad física. En consecuencia, la respuesta penal exige interpretación restrictiva, "
        f"prueba individualizada, daño concreto, proporcionalidad, mínima intervención y control de convencionalidad "
        f"reforzado con enfoque intercultural."
    )

def to_excel_bytes(nombre_hojas: dict) -> bytes:
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        for name, df in nombre_hojas.items():
            safe = str(name)[:31]
            df.to_excel(writer, index=False, sheet_name=safe)
    bio.seek(0)
    return bio.getvalue()

def generar_word_caso(hechos: dict, sub: pd.DataFrame, traz: pd.DataFrame, res: dict) -> bytes | None:
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("Informe de caso MCPI-IRCP-I", 0)

    doc.add_heading("1. Supuesto de hecho", level=1)
    doc.add_paragraph(str(hechos.get("descripcion", "")))

    doc.add_heading("2. Datos principales", level=1)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    tabla.rows[0].cells[0].text = "Campo"
    tabla.rows[0].cells[1].text = "Valor"
    campos = ["pais", "norma", "articulo", "tipo_penal", "sujeto_colectivo", "tipo_protesta", "duracion", "nivel_afectacion", "finalidad", "contexto"]
    for c in campos:
        cells = tabla.add_row().cells
        cells[0].text = c
        cells[1].text = str(hechos.get(c, ""))

    doc.add_heading("3. Matriz de trazabilidad", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["Hecho del caso", "Estado", "Variable afectada", "Impacto metodológico"]):
        t.rows[0].cells[i].text = h
    for _, row in traz.iterrows():
        cells = t.add_row().cells
        for i, h in enumerate(["Hecho del caso", "Estado", "Variable afectada", "Impacto metodológico"]):
            cells[i].text = str(row[h])

    doc.add_heading("4. Subcriterios y valores sugeridos", level=1)
    for variable in sub["Variable"].unique():
        doc.add_heading(variable, level=2)
        s = sub[sub["Variable"] == variable]
        tt = doc.add_table(rows=1, cols=5)
        tt.style = "Table Grid"
        headers = ["Codigo", "Subcriterio", "Peso", "Valor sugerido", "Aporte"]
        for i, h in enumerate(headers):
            tt.rows[0].cells[i].text = h
        for _, row in s.iterrows():
            cells = tt.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row[h])

    doc.add_heading("5. Resultado global", level=1)
    tr = doc.add_table(rows=1, cols=2)
    tr.style = "Table Grid"
    tr.rows[0].cells[0].text = "Indicador"
    tr.rows[0].cells[1].text = "Resultado"
    for k in ["Formal", "Material", "Europeo", "Alexy", "Peso relativo", "CPI", "IRCP-I", "Nivel", "Semáforo"]:
        cells = tr.add_row().cells
        cells[0].text = k
        cells[1].text = str(res.get(k, ""))

    doc.add_heading("6. Conclusión automática", level=1)
    doc.add_paragraph(conclusion_automatica(hechos, res))

    doc.add_heading("7. Reglas de cierre", level=1)
    for regla in reglas_de_cierre(res):
        doc.add_paragraph(regla, style=None)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def aplicar_editor_hechos(hechos: dict, key_prefix: str) -> dict:
    c1, c2, c3 = st.columns(3)
    with c1:
        hechos["pais"] = st.text_input("País", hechos.get("pais", ""), key=f"{key_prefix}_pais")
        hechos["norma"] = st.text_input("Norma", hechos.get("norma", ""), key=f"{key_prefix}_norma")
        hechos["articulo"] = st.text_input("Artículo", hechos.get("articulo", ""), key=f"{key_prefix}_articulo")
    with c2:
        hechos["tipo_penal"] = st.text_input("Tipo penal", hechos.get("tipo_penal", ""), key=f"{key_prefix}_tipo")
        hechos["sujeto_colectivo"] = st.text_input("Sujeto colectivo", hechos.get("sujeto_colectivo", ""), key=f"{key_prefix}_sujeto")
        hechos["tipo_protesta"] = st.text_input("Tipo de protesta", hechos.get("tipo_protesta", ""), key=f"{key_prefix}_protesta")
    with c3:
        hechos["duracion"] = st.selectbox("Duración", ["Minutos", "Varias horas", "Más de un día"], index=["Minutos", "Varias horas", "Más de un día"].index(hechos.get("duracion", "Varias horas")) if hechos.get("duracion", "Varias horas") in ["Minutos", "Varias horas", "Más de un día"] else 1, key=f"{key_prefix}_duracion")
        hechos["nivel_afectacion"] = st.selectbox("Nivel de afectación", ["Sin afectación", "Interrupción parcial de circulación", "Interrupción total", "Afectación grave de servicios esenciales"], index=["Sin afectación", "Interrupción parcial de circulación", "Interrupción total", "Afectación grave de servicios esenciales"].index(hechos.get("nivel_afectacion", "Interrupción parcial de circulación")) if hechos.get("nivel_afectacion", "Interrupción parcial de circulación") in ["Sin afectación", "Interrupción parcial de circulación", "Interrupción total", "Afectación grave de servicios esenciales"] else 1, key=f"{key_prefix}_afectacion")
        hechos["contexto"] = st.text_input("Contexto", hechos.get("contexto", ""), key=f"{key_prefix}_contexto")

    hechos["finalidad"] = st.text_area("Finalidad de la protesta", hechos.get("finalidad", ""), height=70, key=f"{key_prefix}_finalidad")
    hechos["descripcion"] = st.text_area("Descripción completa del supuesto", hechos.get("descripcion", ""), height=140, key=f"{key_prefix}_desc")

    st.subheader("Hechos acreditados / no acreditados")
    opciones = ["Sí", "No", "No consta", "No aplica", "No / deficiente"]
    checks = [
        ("violencia_grave", "Violencia grave"),
        ("danos_personas", "Daños a personas"),
        ("destruccion_bienes", "Destrucción de bienes"),
        ("uso_armas", "Uso de armas"),
        ("organizacion_criminal", "Organización criminal"),
        ("riesgo_concreto_vida_integridad", "Riesgo concreto para vida o integridad"),
        ("dialogo_previo", "Diálogo previo"),
        ("consulta_previa", "Consulta previa"),
        ("alternativas_no_penales", "Alternativas no penales"),
        ("prueba_individualizada", "Prueba individualizada"),
    ]
    for i in range(0, len(checks), 2):
        c1, c2 = st.columns(2)
        for col, (campo, etiqueta) in zip([c1, c2], checks[i:i+2]):
            with col:
                val = hechos.get(campo, "No")
                idx = opciones.index(val) if val in opciones else 1
                hechos[campo] = st.selectbox(etiqueta, opciones, index=idx, key=f"{key_prefix}_{campo}")
    return hechos

def tabla_componentes(res: dict) -> pd.DataFrame:
    return pd.DataFrame([
        {"Componente": "Control formal", "Resultado": res["Formal"], "Peso en CPI": 0.25, "Aporte CPI": round(res["Formal"] * 0.25, 2)},
        {"Componente": "Control material", "Resultado": res["Material"], "Peso en CPI": 0.30, "Aporte CPI": round(res["Material"] * 0.30, 2)},
        {"Componente": "Test europeo", "Resultado": res["Europeo"], "Peso en CPI": 0.25, "Aporte CPI": round(res["Europeo"] * 0.25, 2)},
        {"Componente": "Puntaje Alexy", "Resultado": res["Alexy"], "Peso en CPI": 0.20, "Aporte CPI": round(res["Alexy"] * 0.20, 2)},
    ])

def render_resultados(res: dict, hechos: dict, sub: pd.DataFrame | None = None, traz: pd.DataFrame | None = None):
    st.subheader("Resultado global")
    componentes = tabla_componentes(res)
    st.dataframe(componentes, hide_index=True, use_container_width=True)

    m1, m2, m3 = st.columns(3)
    m1.metric("CPI", res["CPI"])
    m2.metric("IRCP-I", res["IRCP-I"])
    m3.metric("Nivel / semáforo", f"{res['Nivel']} / {res['Semáforo']}")

    st.code(
        f"Formal = ({res['Legalidad']} × 0,60) + ({res['Claridad normativa']} × 0,40) = {res['Formal']}\n"
        f"Material = ({res['Lesividad']} × 0,60) + ({res['Mínima intervención penal']} × 0,40) = {res['Material']}\n"
        f"Europeo = ({res['Idoneidad']} × 0,50) + ({res['Necesidad penal']} × 0,50) = {res['Europeo']}\n"
        f"Alexy = [{res['Peso interés estatal']} ÷ ({res['Peso derechos indígenas']} + {res['Peso interés estatal']})] × 100 = {res['Alexy']}\n"
        f"CPI = ({res['Formal']} × 0,25) + ({res['Material']} × 0,30) + ({res['Europeo']} × 0,25) + ({res['Alexy']} × 0,20) = {res['CPI']}\n"
        f"IRCP-I = 100 − {res['CPI']} = {res['IRCP-I']}",
        language="text"
    )

    st.subheader("Reglas de cierre")
    reglas = reglas_de_cierre(res)
    if reglas:
        for r in reglas:
            st.warning(r)
    else:
        st.success("No se activan reglas críticas de cierre.")

    st.subheader("Conclusión automática")
    st.write(conclusion_automatica(hechos, res))

    chart_df = pd.DataFrame({
        "Componente": ["Formal", "Material", "Europeo", "Alexy", "CPI", "IRCP-I"],
        "Valor": [res["Formal"], res["Material"], res["Europeo"], res["Alexy"], res["CPI"], res["IRCP-I"]],
    })
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.bar(chart_df["Componente"], chart_df["Valor"])
    ax.set_ylim(0, 100)
    ax.set_ylabel("Puntaje")
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    st.pyplot(fig)

    if sub is not None and traz is not None:
        st.download_button(
            "Descargar Excel del caso",
            data=to_excel_bytes({
                "Supuesto": pd.DataFrame([hechos]),
                "Trazabilidad": traz,
                "Subcriterios": sub,
                "Variables": pd.DataFrame([res]),
                "Componentes": componentes,
                "Reglas": pd.DataFrame({"Regla activada": reglas}),
            }),
            file_name="caso_MCPI_IRCP_desde_hechos.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        word = generar_word_caso(hechos, sub, traz, res)
        if word:
            st.download_button(
                "Descargar informe Word del caso",
                data=word,
                file_name="informe_caso_MCPI_IRCP.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

if "base_actual" not in st.session_state:
    st.session_state.base_actual = cargar_base_default()
if "rubrica_actual" not in st.session_state:
    st.session_state.rubrica_actual = cargar_rubrica_default()
if "supuestos_actual" not in st.session_state:
    st.session_state.supuestos_actual = cargar_supuestos_default()
if "hechos_actual" not in st.session_state:
    if not st.session_state.supuestos_actual.empty:
        st.session_state.hechos_actual = st.session_state.supuestos_actual.iloc[0].to_dict()
    else:
        st.session_state.hechos_actual = {}
if "ultimo_resultado" not in st.session_state:
    st.session_state.ultimo_resultado = None
if "ultimo_subcriterios" not in st.session_state:
    st.session_state.ultimo_subcriterios = None
if "ultima_trazabilidad" not in st.session_state:
    st.session_state.ultima_trazabilidad = None

st.title("Matriz MCPI-IRCP-I v7")
st.caption("Versión con supuesto de hecho: hechos acreditados → subcriterios → variables → CPI → IRCP-I.")

menu = st.sidebar.radio(
    "Módulos",
    [
        "Supuesto de hecho",
        "Evaluación desde hechos",
        "Evaluación en matriz",
        "Rúbrica de subcriterios",
        "Base de delitos",
        "Añadir delito",
        "Dashboard comparativo",
        "Análisis de sensibilidad",
        "Exportar informe",
        "Metodología y reglas",
    ],
)

if menu == "Supuesto de hecho":
    st.header("Supuesto de hecho")
    supuestos_df = st.session_state.supuestos_actual

    if not supuestos_df.empty:
        opciones = [f"{i} | {r['nombre_caso']}" for i, r in supuestos_df.iterrows()]
        sel = st.selectbox("Seleccione un supuesto precargado", opciones)
        idx = int(sel.split("|")[0].strip())
        if st.button("Cargar supuesto seleccionado"):
            st.session_state.hechos_actual = supuestos_df.loc[idx].to_dict()
            st.success("Supuesto cargado en la sesión.")

    hechos = dict(st.session_state.hechos_actual)
    hechos = aplicar_editor_hechos(hechos, "supuesto")
    if st.button("Guardar supuesto en la sesión"):
        st.session_state.hechos_actual = hechos
        st.success("Supuesto guardado en sesión.")

    st.download_button(
        "Descargar supuesto actual CSV",
        data=pd.DataFrame([hechos]).to_csv(index=False).encode("utf-8-sig"),
        file_name="supuesto_hecho_actual.csv",
        mime="text/csv",
    )

elif menu == "Evaluación desde hechos":
    st.header("Evaluación desde hechos")
    st.info("La app convierte los hechos acreditados en valores sugeridos. Puede editar esos valores antes de calcular el índice.")
    hechos = dict(st.session_state.hechos_actual)

    with st.expander("Ver / editar supuesto de hecho", expanded=False):
        hechos = aplicar_editor_hechos(hechos, "eval")
        if st.button("Actualizar supuesto para esta evaluación"):
            st.session_state.hechos_actual = hechos
            st.success("Supuesto actualizado.")

    sub, traz, valores = scoring_desde_hechos(hechos, st.session_state.base_actual)

    st.subheader("Matriz de trazabilidad hecho → variable → puntaje")
    st.dataframe(traz, hide_index=True, use_container_width=True)

    st.subheader("Subcriterios sugeridos desde los hechos")
    edited = st.data_editor(
        sub,
        hide_index=True,
        use_container_width=True,
        column_config={
            "Peso": st.column_config.NumberColumn("Peso", disabled=True, format="%.2f"),
            "Valor sugerido": st.column_config.NumberColumn("Valor 0-100", min_value=0, max_value=100, step=1, format="%.0f"),
            "Aporte": st.column_config.NumberColumn("Aporte", disabled=True, format="%.2f"),
            "Explicación metodológica": st.column_config.TextColumn(width="large"),
        },
        disabled=["Modulo", "Variable", "Codigo", "Subcriterio", "Peso", "Aporte", "Explicación metodológica"],
    )
    edited["Valor sugerido"] = pd.to_numeric(edited["Valor sugerido"], errors="coerce").fillna(0).clip(0, 100)
    edited["Peso"] = pd.to_numeric(edited["Peso"], errors="coerce").fillna(0)
    edited["Aporte"] = (edited["Peso"] * edited["Valor sugerido"]).round(2)

    variables = edited.groupby("Variable", as_index=False)["Aporte"].sum().rename(columns={"Aporte": "Valor"})
    st.subheader("Variables calculadas desde subcriterios")
    st.dataframe(variables, hide_index=True, use_container_width=True)

    valores = {r["Variable"]: float(r["Valor"]) for _, r in variables.iterrows()}
    res = calcular_indice_desde_variables(valores)

    st.session_state.ultimo_resultado = res
    st.session_state.ultimo_subcriterios = edited
    st.session_state.ultima_trazabilidad = traz

    render_resultados(res, hechos, edited, traz)

elif menu == "Evaluación en matriz":
    st.header("Evaluación en matriz")
    base = st.session_state.base_actual
    if base.empty:
        st.warning("No hay base de delitos.")
        st.stop()

    opciones = [f"{i} | {r['País']} | {r['Tipo penal']} | {r['Artículo']}" for i, r in base.iterrows()]
    default_idx = next((i for i, x in enumerate(opciones) if "194" in x), 0)
    seleccion = st.selectbox("Seleccione delito/tipo penal", opciones, index=default_idx)
    idx = int(seleccion.split("|")[0].strip())
    row = base.loc[idx].to_dict()

    st.subheader("Datos del tipo penal")
    st.write(f"**{row['País']} · {row['Artículo']} · {row['Tipo penal']}**")
    st.write(row.get("Contexto típico de protesta", ""))

    data = pd.DataFrame([
        {"Variable": "Legalidad", "Puntaje": row["Legalidad"]},
        {"Variable": "Claridad normativa", "Puntaje": row["Claridad normativa"]},
        {"Variable": "Lesividad", "Puntaje": row["Lesividad"]},
        {"Variable": "Mínima intervención penal", "Puntaje": row["Mínima intervención penal"]},
        {"Variable": "Idoneidad", "Puntaje": row["Idoneidad"]},
        {"Variable": "Necesidad penal", "Puntaje": row["Necesidad penal"]},
        {"Variable": "Peso derechos indígenas", "Puntaje": row["Peso derechos indígenas"]},
        {"Variable": "Peso interés estatal", "Puntaje": row["Peso interés estatal"]},
    ])
    edited = st.data_editor(
        data,
        hide_index=True,
        use_container_width=True,
        column_config={
            "Variable": st.column_config.TextColumn(disabled=True),
            "Puntaje": st.column_config.NumberColumn(min_value=0, max_value=100, step=1, format="%.0f"),
        },
        disabled=["Variable"],
    )
    valores = {r["Variable"]: float(r["Puntaje"]) for _, r in edited.iterrows()}
    res = calcular_indice_desde_variables(valores)
    hechos = {
        "pais": row["País"], "norma": row["Norma"], "articulo": row["Artículo"], "tipo_penal": row["Tipo penal"],
        "sujeto_colectivo": "No especificado", "tipo_protesta": row["Contexto típico de protesta"],
        "finalidad": "", "descripcion": row["Observación metodológica"]
    }
    st.session_state.ultimo_resultado = res
    render_resultados(res, hechos)

elif menu == "Rúbrica de subcriterios":
    st.header("Rúbrica de subcriterios")
    st.info("Esta tabla define los subcriterios, pesos y explicación metodológica general.")
    rubrica = st.data_editor(
        st.session_state.rubrica_actual,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={
            "Peso_sobre_variable": st.column_config.NumberColumn("Peso", min_value=0.0, max_value=1.0, step=0.01, format="%.2f"),
            "Valor_editable_0_100": st.column_config.NumberColumn("Valor base 0-100", min_value=0, max_value=100, step=1),
            "Aporte_calculado": st.column_config.NumberColumn("Aporte", disabled=True, format="%.2f"),
        },
    )
    rubrica_norm = normalizar_rubrica(rubrica)
    if st.button("Aplicar cambios de rúbrica a la sesión"):
        st.session_state.rubrica_actual = rubrica_norm
        st.success("Rúbrica actualizada en sesión.")
    st.download_button(
        "Descargar rúbrica Excel",
        data=to_excel_bytes({"Subcriterios_App": rubrica_norm}),
        file_name="subcriterios_MCPI_IRCP_app_unico.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

elif menu == "Base de delitos":
    st.header("Base de delitos")
    edited = st.data_editor(
        st.session_state.base_actual,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={c: st.column_config.NumberColumn(c, min_value=0, max_value=100, step=1) for c in NUMERIC_COLS},
    )
    edited_norm = normalizar_base(edited)
    if st.button("Aplicar cambios a la base en sesión"):
        st.session_state.base_actual = edited_norm
        st.success("Base actualizada.")
    st.download_button(
        "Descargar base CSV",
        data=edited_norm.to_csv(index=False).encode("utf-8-sig"),
        file_name="tipos_penales_base_ampliada.csv",
        mime="text/csv",
    )
    st.download_button(
        "Descargar base Excel",
        data=to_excel_bytes({"Base_delitos": edited_norm, "Resultados": base_resultados(edited_norm)}),
        file_name="tipos_penales_base_ampliada.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

elif menu == "Añadir delito":
    st.header("Añadir nuevo delito/tipo penal")
    with st.form("nuevo_delito"):
        c1, c2, c3 = st.columns(3)
        with c1:
            pais = st.selectbox("País", ["Argentina", "Ecuador", "Otro"])
            norma = st.text_input("Norma", "Código Penal" if pais == "Argentina" else "COIP")
        with c2:
            articulo = st.text_input("Artículo", "")
            categoria = st.text_input("Categoría", "Orden público")
        with c3:
            estado = st.selectbox("Estado", ["Ampliado preliminar", "Base tesis", "Validado por investigador"])
            tipo = st.text_input("Tipo penal", "")
        contexto = st.text_area("Contexto típico de protesta", "")
        obs = st.text_area("Observación metodológica", "")
        st.subheader("Puntajes iniciales")
        p1, p2, p3, p4 = st.columns(4)
        with p1:
            legalidad = st.number_input("Legalidad", 0, 100, 50)
            claridad = st.number_input("Claridad normativa", 0, 100, 45)
        with p2:
            lesividad = st.number_input("Lesividad", 0, 100, 45)
            minima = st.number_input("Mínima intervención penal", 0, 100, 35)
        with p3:
            idoneidad = st.number_input("Idoneidad", 0, 100, 50)
            necesidad = st.number_input("Necesidad penal", 0, 100, 35)
        with p4:
            pdi = st.number_input("Peso derechos indígenas", 0, 100, 40)
            pie = st.number_input("Peso interés estatal", 0, 100, 24)
        submitted = st.form_submit_button("Agregar a la base")
    if submitted:
        nueva = pd.DataFrame([{
            "País": pais, "Norma": norma, "Artículo": articulo, "Tipo penal": tipo, "Categoría": categoria,
            "Estado": estado, "Contexto típico de protesta": contexto, "Legalidad": legalidad, "Claridad normativa": claridad,
            "Lesividad": lesividad, "Mínima intervención penal": minima, "Idoneidad": idoneidad, "Necesidad penal": necesidad,
            "Peso derechos indígenas": pdi, "Peso interés estatal": pie, "Observación metodológica": obs
        }])
        st.session_state.base_actual = normalizar_base(pd.concat([st.session_state.base_actual, nueva], ignore_index=True))
        st.success("Delito agregado a la base de la sesión.")

elif menu == "Dashboard comparativo":
    st.header("Dashboard comparativo")
    resultados = base_resultados(st.session_state.base_actual)
    if resultados.empty:
        st.warning("No hay datos.")
        st.stop()
    st.dataframe(resultados, hide_index=True, use_container_width=True)
    c1, c2, c3 = st.columns(3)
    c1.metric("Delitos en base", len(resultados))
    c2.metric("IRCP-I promedio", round(resultados["IRCP-I"].mean(), 2))
    c3.metric("Mayor riesgo", resultados.loc[resultados["IRCP-I"].idxmax(), "Tipo penal"])

    plot = resultados.sort_values("IRCP-I")
    fig, ax = plt.subplots(figsize=(11, max(5, len(plot) * 0.28)))
    ax.barh(plot["País"] + " - " + plot["Tipo penal"], plot["IRCP-I"])
    ax.set_xlim(0, 100)
    ax.set_xlabel("IRCP-I")
    ax.grid(axis="x", linestyle="--", alpha=0.4)
    st.pyplot(fig)

elif menu == "Análisis de sensibilidad":
    st.header("Análisis de sensibilidad")
    if st.session_state.ultimo_resultado is None:
        hechos = dict(st.session_state.hechos_actual)
        sub, traz, valores = scoring_desde_hechos(hechos, st.session_state.base_actual)
        res = calcular_indice_desde_variables(valores)
    else:
        res = st.session_state.ultimo_resultado

    escenarios = pd.DataFrame([
        {"Escenario": "Modelo base", "Formal": 0.25, "Material": 0.30, "Europeo": 0.25, "Alexy": 0.20},
        {"Escenario": "Pesos iguales", "Formal": 0.25, "Material": 0.25, "Europeo": 0.25, "Alexy": 0.25},
        {"Escenario": "Mayor peso material", "Formal": 0.20, "Material": 0.35, "Europeo": 0.25, "Alexy": 0.20},
        {"Escenario": "Mayor peso intercultural", "Formal": 0.20, "Material": 0.25, "Europeo": 0.25, "Alexy": 0.30},
    ])

    rows = []
    for _, e in escenarios.iterrows():
        cpi = round(res["Formal"] * e["Formal"] + res["Material"] * e["Material"] + res["Europeo"] * e["Europeo"] + res["Alexy"] * e["Alexy"], 2)
        ircp = round(100 - cpi, 2)
        nivel, sem = nivel_riesgo(ircp)
        rows.append({
            "Escenario": e["Escenario"],
            "Peso Formal": e["Formal"],
            "Peso Material": e["Material"],
            "Peso Europeo": e["Europeo"],
            "Peso Alexy": e["Alexy"],
            "CPI": cpi,
            "IRCP-I": ircp,
            "Nivel": nivel,
            "Semáforo": sem,
        })
    sensibilidad = pd.DataFrame(rows)
    st.dataframe(sensibilidad, hide_index=True, use_container_width=True)

    fig, ax = plt.subplots(figsize=(9, 4))
    ax.bar(sensibilidad["Escenario"], sensibilidad["IRCP-I"])
    ax.set_ylim(0, 100)
    ax.set_ylabel("IRCP-I")
    ax.tick_params(axis="x", rotation=20)
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    st.pyplot(fig)

    st.download_button(
        "Descargar análisis de sensibilidad Excel",
        data=to_excel_bytes({"Sensibilidad": sensibilidad}),
        file_name="analisis_sensibilidad_MCPI_IRCP.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

elif menu == "Exportar informe":
    st.header("Exportar informe")
    if st.session_state.ultimo_resultado is None:
        st.info("Primero ejecute una evaluación desde hechos. Se generará el informe del supuesto actualmente cargado.")
        hechos = dict(st.session_state.hechos_actual)
        sub, traz, valores = scoring_desde_hechos(hechos, st.session_state.base_actual)
        res = calcular_indice_desde_variables(valores)
    else:
        hechos = dict(st.session_state.hechos_actual)
        sub = st.session_state.ultimo_subcriterios
        traz = st.session_state.ultima_trazabilidad
        res = st.session_state.ultimo_resultado

    st.write(conclusion_automatica(hechos, res))
    word = generar_word_caso(hechos, sub, traz, res)
    if word:
        st.download_button(
            "Descargar informe Word del caso",
            data=word,
            file_name="informe_caso_MCPI_IRCP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    st.download_button(
        "Descargar Excel completo del caso",
        data=to_excel_bytes({
            "Supuesto": pd.DataFrame([hechos]),
            "Trazabilidad": traz,
            "Subcriterios": sub,
            "Resultado": pd.DataFrame([res]),
            "Componentes": tabla_componentes(res),
        }),
        file_name="caso_MCPI_IRCP_completo.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

else:
    st.header("Metodología y reglas")
    st.markdown("""
### Flujo metodológico de la app v7

**Hechos acreditados → subcriterios → variables → componentes → CPI → IRCP-I**

### Fórmulas

**Control formal** = (Legalidad × 0,60) + (Claridad normativa × 0,40)

**Control material** = (Lesividad × 0,60) + (Mínima intervención penal × 0,40)

**Test europeo** = (Idoneidad × 0,50) + (Necesidad penal × 0,50)

**Puntaje Alexy** = Peso interés estatal / (Peso derechos indígenas + Peso interés estatal) × 100

**CPI** = (Formal × 0,25) + (Material × 0,30) + (Europeo × 0,25) + (Alexy × 0,20)

**IRCP-I** = 100 − CPI

### Regla de lectura

- A mayor CPI, mayor compatibilidad penal intercultural.
- A mayor IRCP-I, mayor riesgo de criminalización penal intercultural.

### Escala de riesgo

| IRCP-I | Nivel |
|---:|---|
| 0–34 | Riesgo bajo |
| 35–49 | Riesgo medio |
| 50–69 | Riesgo medio-alto |
| 70–100 | Riesgo alto |

### Persistencia

Los cambios en Streamlit quedan en la sesión. Para conservarlos, descargue CSV/Excel y reemplace el archivo correspondiente en GitHub.
""")

st.caption("Herramienta académica orientativa. No sustituye análisis jurídico definitivo ni control judicial del caso concreto.")
